from __future__ import annotations
"""
ingestion.py - EduBot Document Ingestion Pipeline

Supports:
- PDF, DOCX, TXT, HTML, CSV, Markdown, JSON, Excel, SQL dumps, Images OCR
- Website crawling
- Website quick links extraction
- Website PDF link discovery
- Website PDF content ingestion
"""

from db import add_chunks, collection, normalize_metadata
from embeddings import encode_texts
from rag.bm25_index import rebuild_bm25_index

import os
import sys
import io
import re
import uuid
import time
import json
import hashlib
import warnings
from urllib.parse import parse_qs, unquote, urljoin, urlparse
from urllib.robotparser import RobotFileParser
from pathlib import Path

import requests
import trafilatura
from bs4 import BeautifulSoup

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import (
    TextLoader,
    Docx2txtLoader,
    UnstructuredHTMLLoader,
)

try:
    from PIL import Image, ImageOps, ImageFilter
    # Decompression-bomb guard: cap the pixel count Pillow will decode. A small
    # crafted image header can claim enormous dimensions and exhaust memory on
    # decode. Default ~64 MP (covers an A3 page at 300 DPI); override via env.
    try:
        Image.MAX_IMAGE_PIXELS = int(os.getenv("MAX_IMAGE_PIXELS", str(64_000_000)))
    except Exception:
        pass
except Exception:
    Image = None
    ImageOps = None
    ImageFilter = None

try:
    import pytesseract
except Exception:
    pytesseract = None

import pdfplumber
import pandas as pd
import threading

from rag.freshness import (
    extract_year_from_text,
    normalize_source_type,
    parse_document_date,
    parse_document_year,
)
from rag.text_utils import normalize_ligatures

# =============================================================================
# CRAWL JOB MANAGEMENT & HELPER FUNCTIONS
# =============================================================================

crawl_jobs = {}
crawl_jobs_lock = threading.Lock()

# Local JSON persistence for table-free/schema-independent operation
CRAWL_JOBS_FILE = Path(__file__).resolve().parent / "data" / "crawl_jobs.json"

def load_crawl_jobs_from_file():
    global crawl_jobs
    try:
        if CRAWL_JOBS_FILE.exists():
            with open(CRAWL_JOBS_FILE, "r") as f:
                data = json.load(f)
                if isinstance(data, dict):
                    crawl_jobs.update(data)
                    print(f"[Crawl Jobs] Loaded {len(data)} crawl jobs from local storage.")
    except Exception as e:
        print(f"[Crawl Jobs Load] Warning: failed to load from file: {e}")

def save_crawl_jobs_to_file():
    try:
        # Ensure directory exists
        CRAWL_JOBS_FILE.parent.mkdir(parents=True, exist_ok=True)
        with crawl_jobs_lock:
            jobs_to_save = dict(crawl_jobs)
        with open(CRAWL_JOBS_FILE, "w") as f:
            json.dump(jobs_to_save, f, indent=2)
    except Exception as e:
        print(f"[Crawl Jobs Save] Warning: failed to save to file: {e}")

# Load immediately on startup
load_crawl_jobs_from_file()

def sync_crawl_job_to_supabase(job_id: str):
    try:
        import os
        import requests
        supabase_url = os.getenv("SUPABASE_URL", "")
        supabase_key = os.getenv("SUPABASE_SERVICE_ROLE_KEY", "")
        if not supabase_url or not supabase_key:
            return
            
        with crawl_jobs_lock:
            job = crawl_jobs.get(job_id)
            if not job:
                return
            job_data = dict(job)
            
        # Map values to Supabase fields
        payload = {
            "job_id": job_id,
            "url": job_data.get("url", "") or job_data.get("current_url", "") or "",
            "status": job_data.get("status", "pending"),
            "current_stage": job_data.get("current_stage", "queued"),
            "pages_found": job_data.get("pages_found", 0),
            "pages_processed": job_data.get("pages_processed", 0),
            "pages_skipped": len(job_data.get("skipped_urls", [])),
            "pages_failed": job_data.get("pages_failed", 0),
            "pdfs_found": job_data.get("pdfs_found", 0),
            "pdfs_processed": job_data.get("pdfs_processed", 0),
            "documents_found": job_data.get("documents_found", 0),
            "documents_processed": job_data.get("documents_processed", 0),
            "chunks_created": job_data.get("chunks_created", 0),
            "embeddings_generated": job_data.get("embeddings_generated", 0),
            "skipped_urls": job_data.get("skipped_urls", []),
            "errors": job_data.get("errors", []),
            "started_at": job_data.get("started_at"),
            "finished_at": job_data.get("finished_at"),
            "last_crawl_timestamp": job_data.get("last_crawl_timestamp", time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()))
        }
        
        headers = {
            "apikey": supabase_key,
            "Authorization": f"Bearer {supabase_key}",
            "Content-Type": "application/json",
            "Prefer": "resolution=merge-duplicates"
        }
        
        url = f"{supabase_url.rstrip('/')}/rest/v1/crawl_jobs"
        requests.post(url, json=payload, headers=headers, timeout=5)
    except Exception as e:
        print(f"[Supabase Sync] Warning: failed to sync crawl job {job_id}: {e}")


def create_crawl_job() -> str:
    job_id = str(uuid.uuid4())
    with crawl_jobs_lock:
        crawl_jobs[job_id] = {
            "job_id": job_id,
            "url": "",
            "status": "queued",  # queued / pending / crawling / processing / chunking / embedding / completed / failed / cancelled / paused
            "current_url": "",
            "current_type": "",  # page / pdf / document
            "current_stage": "queued",  # queued / fetching / extracting / downloading / ingesting / skipped / completed
            "pages_found": 0,
            "pages_processed": 0,
            "pages_skipped": 0,
            "pages_failed": 0,
            "pdfs_found": 0,
            "pdfs_processed": 0,
            "documents_found": 0,
            "documents_processed": 0,
            "chunks_created": 0,
            "embeddings_generated": 0,
            "skipped_urls": [],
            "errors": [],
            "started_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
            "finished_at": None,
            "last_crawl_timestamp": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
            # Control flags
            "skip_current_page": False,
            "skip_current_document": False,
            "pause": False,
            "cancel": False
        }
    save_crawl_jobs_to_file()
    return job_id


last_sync_times = {}
last_sync_lock = threading.Lock()


def _should_sync_immediately(status: str | None, current_stage: str | None) -> bool:
    if status in {"completed", "failed", "cancelled", "started", "pending"}:
        return True
    if current_stage in {"completed", "failed", "cancelled"}:
        return True
    return False


def _persist_and_sync_job_background(job_id: str):
    try:
        sync_crawl_job_to_supabase(job_id)
    except Exception as e:
        print(f"[Supabase Sync Background] Warning: {e}")


def _persist_and_sync_job(job_id: str, force: bool = False):
    now = time.time()
    should_sync = force

    with crawl_jobs_lock:
        job = crawl_jobs.get(job_id)
        if job:
            status = job.get("status")
            current_stage = job.get("current_stage")
            if _should_sync_immediately(status, current_stage):
                should_sync = True

    if not should_sync:
        with last_sync_lock:
            last_sync = last_sync_times.get(job_id, 0.0)
            if now - last_sync >= 5.0:
                last_sync_times[job_id] = now
                should_sync = True

    if should_sync:
        save_crawl_jobs_to_file()
        threading.Thread(
            target=_persist_and_sync_job_background,
            args=(job_id,),
            daemon=True
        ).start()


def update_crawl_job(job_id: str | None, **fields):
    if not job_id:
        return
    with crawl_jobs_lock:
        if job_id in crawl_jobs:
            job = crawl_jobs[job_id]
            for k, v in fields.items():
                if k == "errors":
                    job["errors"].append(v)
                elif k == "skipped_urls":
                    job["skipped_urls"].append(v)
                else:
                    job[k] = v
            job["last_crawl_timestamp"] = time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
            
    _persist_and_sync_job(job_id)


def increment_crawl_job(job_id: str | None, metrics: dict[str, int], current_stage: str | None = None):
    if not job_id:
        return
    with crawl_jobs_lock:
        if job_id in crawl_jobs:
            job = crawl_jobs[job_id]
            for k, v in metrics.items():
                job[k] = job.get(k, 0) + v
            if current_stage:
                job["current_stage"] = current_stage
            job["last_crawl_timestamp"] = time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())

    _persist_and_sync_job(job_id)



def get_all_crawl_jobs() -> list[dict]:
    with crawl_jobs_lock:
        jobs = []
        for job in crawl_jobs.values():
            job_copy = dict(job)
            # Add PDF fallback names for compatibility
            job_copy["PDFs found"] = job_copy.get("pdfs_found", 0)
            job_copy["PDFs processed"] = job_copy.get("pdfs_processed", 0)
            jobs.append(job_copy)
        
        # Sort by started_at desc
        try:
            jobs.sort(key=lambda x: x.get("started_at", ""), reverse=True)
        except Exception:
            pass
        return jobs


def delete_crawl_job(job_id: str | None) -> bool:
    if not job_id:
        return False
    # Only the in-memory mutation needs the lock. save_crawl_jobs_to_file() also
    # acquires crawl_jobs_lock, and threading.Lock is non-reentrant, so persisting
    # while still holding the lock here would deadlock the request (the DELETE would
    # hang forever). Mirror update_crawl_job(): mutate under the lock, persist outside.
    with crawl_jobs_lock:
        if job_id not in crawl_jobs:
            return False
        del crawl_jobs[job_id]

    save_crawl_jobs_to_file()

    # Sync delete to Supabase if config is present (best effort)
    try:
        import os
        import requests
        supabase_url = os.getenv("SUPABASE_URL", "")
        supabase_key = os.getenv("SUPABASE_SERVICE_ROLE_KEY", "")
        if supabase_url and supabase_key:
            headers = {
                "apikey": supabase_key,
                "Authorization": f"Bearer {supabase_key}",
                "Content-Type": "application/json"
            }
            url = f"{supabase_url.rstrip('/')}/rest/v1/crawl_jobs?job_id=eq.{job_id}"
            requests.delete(url, headers=headers, timeout=5)
    except Exception as e:
        print(f"[Supabase Sync] Warning: failed to delete crawl job {job_id} from Supabase: {e}")

    return True

def get_crawl_job(job_id: str | None) -> dict | None:
    if not job_id:
        return None
    with crawl_jobs_lock:
        job = crawl_jobs.get(job_id)
        if not job:
            return None
        job_copy = dict(job)
        job_copy["PDFs found"] = job_copy.get("pdfs_found", 0)
        job_copy["PDFs processed"] = job_copy.get("pdfs_processed", 0)
        return job_copy

def control_crawl_job(job_id: str | None, action: str) -> bool:
    if not job_id:
        return False
    with crawl_jobs_lock:
        if job_id not in crawl_jobs:
            return False
        job = crawl_jobs[job_id]
        if action == "skip_current_page":
            job["skip_current_page"] = True
        elif action == "skip_current_document":
            job["skip_current_document"] = True
        elif action == "pause":
            job["pause"] = True
            if job["status"] == "running":
                job["status"] = "paused"
        elif action == "resume":
            job["pause"] = False
            job["status"] = "running"
        elif action == "cancel":
            job["cancel"] = True
            job["status"] = "cancelled"
            job["finished_at"] = time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
        return True

def should_cancel(job_id: str | None) -> bool:
    if not job_id:
        return False
    with crawl_jobs_lock:
        job = crawl_jobs.get(job_id)
        if job and (job.get("cancel") or job.get("status") == "cancelled"):
            return True
    return False

def check_crawl_cancelled(job_id: str | None):
    if should_cancel(job_id):
        raise ValueError("Crawl cancelled by admin during ingestion.")

def wait_if_paused(job_id: str | None):
    if not job_id:
        return
    while True:
        with crawl_jobs_lock:
            job = crawl_jobs.get(job_id)
            if not job:
                return
            if job.get("cancel") or job.get("status") == "cancelled":
                raise ValueError("Crawl cancelled by admin during ingestion.")
            if not job.get("pause"):
                if job.get("status") == "paused":
                    job["status"] = "running"
                return
        time.sleep(0.5)

def should_skip_current_page(job_id: str | None) -> bool:
    if not job_id:
        return False
    with crawl_jobs_lock:
        job = crawl_jobs.get(job_id)
        if job and job.get("skip_current_page"):
            job["skip_current_page"] = False
            return True
    return False

def should_skip_current_document(job_id: str | None) -> bool:
    if not job_id:
        return False
    with crawl_jobs_lock:
        job = crawl_jobs.get(job_id)
        if job and job.get("skip_current_document"):
            job["skip_current_document"] = False
            return True
    return False

def run_crawl_background(
    job_id: str,
    url: str,
    department: str,
    document_type: str,
    year: str,
    max_pages: int,
    include_pdfs: bool,
    max_pdfs: int,
    same_domain_only: bool,
    max_depth: int = 3,
):
    try:
        update_crawl_job(
            job_id,
            url=url,
            status="crawling",
            started_at=time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
            max_pages=max_pages,
            max_pdfs=max_pdfs,
            max_depth=max_depth,
        )
        result = ingest_website(
            url=url,
            department=department,
            document_type=document_type,
            year=year,
            scope="official",
            max_pages=max_pages,
            include_pdfs=include_pdfs,
            max_pdfs=max_pdfs,
            same_domain_only=same_domain_only,
            job_id=job_id,
            max_depth=max_depth,
        )
        if should_cancel(job_id):
            update_crawl_job(
                job_id,
                status="cancelled",
                finished_at=time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
            )
        else:
            update_crawl_job(
                job_id,
                status="completed",
                finished_at=time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
            )
    except Exception as e:
        error_msg = str(e)
        if "Crawl cancelled" in error_msg or should_cancel(job_id):
            update_crawl_job(
                job_id,
                status="cancelled",
                errors="Crawl cancelled by admin during ingestion.",
                finished_at=time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
            )
        else:
            update_crawl_job(
                job_id,
                status="failed",
                errors=f"Serious error: {error_msg}",
                finished_at=time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
            )


# =============================================================================
# CONFIG
# =============================================================================

BASE_DIR = Path(__file__).resolve().parent
UPLOAD_DIR = BASE_DIR / "data" / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".html",
    ".htm",
    ".csv",
    ".md",
    ".markdown",
    ".json",
    ".xlsx",
    ".xls",
    ".sql",
    ".dump",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
}

DOCUMENT_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".doc",
    ".xlsx",
    ".xls",
    ".csv",
    ".ppt",
    ".pptx",
    ".txt",
    ".md",
    ".json",
}

# Constants imported from rag.config for configuration consistency
from rag.config import (
    MAX_CHARS_PER_CHUNK,
    MAX_CHARS_PER_LIST_CHUNK,
    MIN_CHUNK_WORDS,
    CHUNK_MARGIN_FACTOR,
    CHUNK_OVERLAP,
)
PDF_OCR_FALLBACK_MIN_WORDS = 20

REQUEST_TIMEOUT_SECONDS = 25

WEBSITE_BLOCKED_EXTENSIONS = (
    ".svg",
    ".ico",
    ".mp4",
    ".mp3",
    ".avi",
    ".mov",
    ".mkv",
    ".zip",
    ".rar",
    ".7z",
    ".css",
    ".js",
    ".woff",
    ".woff2",
    ".ttf",
    ".otf",
    ".eot",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".gif",
    ".bmp",
    ".tiff",
)

CRAWLER_HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.9",
    "Accept-Encoding": "gzip, deflate, br",
    "Connection": "keep-alive",
    "Upgrade-Insecure-Requests": "1",
    "Cache-Control": "max-age=0",
}


def create_robust_session(headers: dict = None) -> requests.Session:
    from urllib3.util import Retry
    from requests.adapters import HTTPAdapter
    session = requests.Session()
    if headers:
        session.headers.update(headers)
    
    # Retry on connection/read errors, connection drops, rate limiting (429), and server issues
    retries = Retry(
        total=3,
        backoff_factor=1,  # wait 1s, 2s, 4s...
        status_forcelist=[429, 500, 502, 503, 504],
        raise_on_status=False,
    )
    adapter = HTTPAdapter(max_retries=retries)
    session.mount("http://", adapter)
    session.mount("https://", adapter)
    return session


# Maximum bytes accepted for any single crawled page/document. Caps memory use
# and defends against oversized files and decompression bombs (iter_content yields
# already-decompressed bytes, so the cap applies to the DECODED size). Override
# via MAX_DOWNLOAD_BYTES; default 50 MB comfortably covers college prospectuses.
MAX_DOWNLOAD_BYTES = int(os.getenv("MAX_DOWNLOAD_BYTES", str(50 * 1024 * 1024)))

# Memory-safety cap on per-PDF page processing (rendering + OCR is expensive). A
# crafted/huge PDF beyond this many pages has the remainder skipped. Override via
# MAX_PDF_PAGES; default 2000 covers any real college document.
MAX_PDF_PAGES = int(os.getenv("MAX_PDF_PAGES", "2000"))

# ZIP-expansion (decompression-bomb) protection for ZIP-based Office formats
# (DOCX, XLSX, PPTX). Limits are generous for real college documents but reject
# crafted archives that explode on decompression. All overridable via env.
MAX_ARCHIVE_MEMBERS = int(os.getenv("MAX_ARCHIVE_MEMBERS", "10000"))
MAX_ARCHIVE_MEMBER_BYTES = int(os.getenv("MAX_ARCHIVE_MEMBER_BYTES", str(100 * 1024 * 1024)))
MAX_ARCHIVE_TOTAL_BYTES = int(os.getenv("MAX_ARCHIVE_TOTAL_BYTES", str(500 * 1024 * 1024)))
MAX_ARCHIVE_COMPRESSION_RATIO = float(os.getenv("MAX_ARCHIVE_COMPRESSION_RATIO", "200"))
# Below this uncompressed size the ratio check is skipped: small, legitimately
# repetitive Office XML can have a high ratio yet be harmless.
_ARCHIVE_RATIO_MIN_BYTES = 10 * 1024 * 1024


def _read_response_with_size_cap(response, url: str):
    """Buffer a streamed response body up to MAX_DOWNLOAD_BYTES, raising when the
    advertised or actual (decompressed) size exceeds the cap. Sets ``_content`` so
    callers can use ``response.content``/``response.text`` normally afterwards."""
    advertised = response.headers.get("Content-Length")
    if advertised and advertised.isdigit() and int(advertised) > MAX_DOWNLOAD_BYTES:
        response.close()
        raise ValueError(
            f"Refusing {url!r}: Content-Length {int(advertised)} exceeds "
            f"{MAX_DOWNLOAD_BYTES}-byte download cap."
        )
    total = 0
    parts: list[bytes] = []
    for part in response.iter_content(chunk_size=65536):
        if not part:
            continue
        total += len(part)
        if total > MAX_DOWNLOAD_BYTES:
            response.close()
            raise ValueError(
                f"Refusing {url!r}: download exceeded {MAX_DOWNLOAD_BYTES}-byte cap."
            )
        parts.append(part)
    response._content = b"".join(parts)
    response._content_consumed = True
    return response


def fetch_with_ssl_fallback(session, url: str, **kwargs):
    """GET ``url``, retrying once with TLS verification disabled if the host
    presents an invalid/mismatched certificate.

    Many college sites have a certificate that covers the apex domain
    (anthonys.ac.in) but not the ``www`` host, so a normal verified request
    raises ``SSLError`` and the whole crawl aborts. On that specific failure we
    retry with ``verify=False`` and log a clear warning so the insecure fetch is
    visible. Any other error propagates unchanged.

    The body is streamed and size-capped (see ``_read_response_with_size_cap``)
    so a single oversized file or decompression bomb cannot exhaust memory.
    """
    # Stream so the size cap can abort before the whole body is buffered.
    kwargs.setdefault("stream", True)
    try:
        response = session.get(url, **kwargs)
    except requests.exceptions.SSLError as ssl_err:
        print(
            f"[WEBSITE] SSL verification failed for {url}: {ssl_err}. "
            f"Retrying with certificate verification disabled (insecure)."
        )
        kwargs["verify"] = False
        # Silence the expected InsecureRequestWarning for this deliberate retry.
        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            response = session.get(url, **kwargs)
    return _read_response_with_size_cap(response, url)



# =============================================================================
# CLEANING + CHUNKING
# =============================================================================

PDF_MIXED_CASE_FIXES = {
    "admission": "Admission",
    "admissions": "Admissions",
    "affairs": "Affairs",
    "alumni": "Alumni",
    "anti": "Anti",
    "award": "Award",
    "canteen": "Canteen",
    "cell": "Cell",
    "committee": "Committee",
    "committees": "Committees",
    "department": "Department",
    "departments": "Departments",
    "examination": "Examination",
    "fee": "Fee",
    "fees": "Fees",
    "graduate": "Graduate",
    "monitoring": "Monitoring",
    "post": "Post",
    "ragging": "Ragging",
    "structure": "Structure",
    "under": "Under",
}


def repair_pdf_mixed_case_terms(text: str) -> str:
    if not text:
        return ""

    def replace_word(match: re.Match) -> str:
        word = match.group(0)

        if not (any(c.islower() for c in word) and any(c.isupper() for c in word)):
            return word

        return PDF_MIXED_CASE_FIXES.get(word.lower(), word)

    text = re.sub(r"\b[A-Za-z]+\b", replace_word, text)
    text = re.sub(r"\bAnti\s*-\s*ragging\b", "Anti-Ragging", text, flags=re.IGNORECASE)

    return text


def clean_loaded_text(text: str) -> str:
    if not text:
        return ""

    text = str(text)
    text = text.replace("\x00", " ")
    text = normalize_ligatures(text)
    text = repair_pdf_mixed_case_terms(text)
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    text = re.sub(
        r"^\s*(?:[-–]\s*)?\d+\s*(?:[-–]\s*)?$",
        "",
        text,
        flags=re.MULTILINE,
    )

    text = re.sub(
        r"^\s*Page\s+\d+(?:\s+of\s+\d+)?\s*$",
        "",
        text,
        flags=re.IGNORECASE | re.MULTILINE,
    )

    text = re.sub(
        r"^\s*(?:st\.?\s*anthony\'?s?\s*college|shillong|prospectus|handbook|syllabus|course structure)\s*$",
        "",
        text,
        flags=re.IGNORECASE | re.MULTILINE,
    )

    text = re.sub(
        r"^\s*\d+\s*\|\s*P\s*a\s*g\s*e\s*$",
        "",
        text,
        flags=re.IGNORECASE | re.MULTILINE,
    )

    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" *\n *", "\n", text)

    return text.strip()


def compute_text_hash(text: str) -> str:
    return hashlib.sha256(str(text or "").strip().encode("utf-8")).hexdigest()[:24]


# Near-duplicate detection (SimHash). Two chunks are "near-duplicates" only when
# their 64-bit SimHash differs in <= this many bits.
SIMHASH_NEAR_DUP_MAX_HAMMING = 3


def _simhash(text: str, bits: int = 64) -> int:
    """64-bit SimHash over word tokens — small Hamming distance ⇒ near-identical text."""
    tokens = re.findall(r"[a-z0-9]+", str(text or "").lower())
    if not tokens:
        return 0
    vector = [0] * bits
    for token in tokens:
        h = int.from_bytes(hashlib.md5(token.encode("utf-8")).digest()[:8], "big")
        for i in range(bits):
            vector[i] += 1 if (h >> i) & 1 else -1
    out = 0
    for i in range(bits):
        if vector[i] > 0:
            out |= 1 << i
    return out


def _hamming(a: int, b: int) -> int:
    return bin(a ^ b).count("1")


def _digit_signature(text: str) -> tuple:
    """Sorted multiset of number-strings in the text.

    Used as a guard: two chunks are only merged as near-duplicates when this
    matches, so distinct numeric data (e.g. 'Fee 5000' vs 'Fee 6000', different
    dates/marks) is NEVER collapsed even if the surrounding wording is identical.
    """
    return tuple(sorted(re.findall(r"\d+", str(text or ""))))


def word_count(text: str) -> int:
    return len(str(text or "").split())


def compute_doc_id(filename, source_url="") -> str:
    key = clean_loaded_text(f"{filename or ''}\n{source_url or ''}") or "unknown"
    digest = hashlib.sha256(key.encode("utf-8")).hexdigest()[:24]

    return f"doc_{digest}"


def utc_timestamp() -> str:
    return time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())


def metadata_year_value(*values) -> str:
    year = extract_year_from_text(*values)
    return str(year) if year else ""


def metadata_date_value(meta: dict | None = None, text: str = "") -> str:
    return parse_document_date(meta or {}, text) or ""


def detect_file_type(filepath: str) -> str:
    import mimetypes
    mime_type, _ = mimetypes.guess_type(filepath)
    if mime_type:
        return mime_type
    ext = os.path.splitext(filepath)[1].lower()
    return ext or "unknown"


def detect_chunk_type(text) -> str:
    text = str(text or "").strip()

    if not text:
        return "empty"

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    lower = text.lower()

    if is_markdown_table(text):
        return "table"

    if lines and all(re.match(r"^[-*+]\s+\S+", line) for line in lines[:8]):
        return "list"

    if lines and all(re.match(r"^\d+[\).]\s+\S+", line) for line in lines[:8]):
        return "list"

    if lower.startswith("website links / quick links") or "\n- " in text:
        return "links"

    if "pdf document found on the college website" in lower:
        return "pdf_link"

    if len(lines) == 1 and looks_like_section_heading(lines[0]):
        return "heading"

    return "text"


def detect_table_title(text, section_title="") -> str:
    if not is_markdown_table(text):
        return ""

    section_title = str(section_title or "").strip()

    if section_title and section_title.lower() != "general":
        return section_title

    for line in str(text or "").splitlines():
        stripped = line.strip().strip("|").strip()

        if stripped and not set(stripped.replace("|", "").replace(" ", "")) <= {"-"}:
            return stripped[:120]

    return ""


def build_heading_path(section_title, doc_metadata=None) -> str:
    doc_metadata = doc_metadata or {}
    parts = []

    for key in ("document_title", "title", "sheet_name"):
        value = str(doc_metadata.get(key, "") or "").strip()

        if value:
            parts.append(value)
            break

    section = str(section_title or "").strip()

    if section and section.lower() != "general":
        parts.append(section)

    if not parts:
        return "general"

    deduped = []

    for part in parts:
        if part not in deduped:
            deduped.append(part)

    return " > ".join(deduped)


def find_chunk_char_offsets(full_text, chunk, search_start=0) -> tuple[int, int]:
    full_text = str(full_text or "")
    chunk = str(chunk or "")

    if not full_text or not chunk:
        return -1, -1

    start = full_text.find(chunk, max(int(search_start or 0), 0))

    if start < 0:
        start = full_text.find(chunk)

    if start < 0:
        compact_full = re.sub(r"\s+", " ", full_text)
        compact_chunk = re.sub(r"\s+", " ", chunk).strip()
        compact_start = compact_full.find(compact_chunk)

        if compact_start < 0:
            return -1, -1

        return compact_start, compact_start + len(compact_chunk)

    return start, start + len(chunk)


def is_markdown_table(text: str) -> bool:
    lines = [line.strip() for line in str(text or "").splitlines() if line.strip()]
    return len(lines) >= 2 and any(line.startswith("|") and line.endswith("|") for line in lines)


def looks_like_section_heading(line: str) -> bool:
    line = str(line or "").strip()

    if not line or is_markdown_table(line):
        return False

    words = line.split()

    if len(words) > 14:
        return False

    if re.match(r"^#{1,6}\s+\S+", line):
        return True

    if re.match(r"^(\d+(\.\d+)*|[IVXLC]+|[A-Z])[\).:-]\s+\S+", line, re.IGNORECASE):
        return True

    lower = line.lower()

    heading_keywords = (
        "admission",
        "attendance",
        "committee",
        "contact",
        "course",
        "department",
        "eligibility",
        "examination",
        "faculty",
        "fee",
        "fees",
        "hostel",
        "library",
        "mission",
        "programme",
        "rules",
        "staff",
        "syllabus",
        "vision",
        "notice",
        "news",
        "quick links",
        "downloads",
    )

    if any(kw in lower for kw in heading_keywords):
        return True

    letters = [c for c in line if c.isalpha()]

    if letters and sum(1 for c in letters if c.isupper()) / len(letters) >= 0.75:
        return True

    return False


def detect_toc(text: str) -> bool:
    lines = [line.strip() for line in str(text or "").splitlines() if line.strip()]

    if not lines:
        return False

    sample = "\n".join(lines[:30]).lower()

    if "table of contents" in sample or (lines and lines[0].lower().strip() in {"contents", "index", "table of contents"}):
        return True

    dotted = sum(1 for line in lines[:50] if re.search(r"\.{3,}\s*\d+\s*$", line))
    numbered = sum(
        1
        for line in lines[:50]
        if re.search(r"\s+\d{1,3}\s*$", line) and len(line.split()) <= 10
    )

    return dotted >= 3 or numbered >= 6


def chunk_text(
    text: str,
    max_length: int = MAX_CHARS_PER_CHUNK,
    overlap: int = CHUNK_OVERLAP,
) -> list[str]:
    """
    Structure-aware chunking for EduBot ingestion.
    Respects boundaries of tables, lists, and headings, keeping them intact up to
    CHUNK_MARGIN_FACTOR * limit, and falls back to character-based splitting with overlap when necessary.
    """
    text = clean_loaded_text(text)
    if not text:
        return []

    # Detect if the overall text is predominantly a list or table to adjust limit
    is_list = detect_chunk_type(text) in ("list", "links")
    limit = MAX_CHARS_PER_LIST_CHUNK if is_list else max_length

    # 1. Split the text into structural blocks
    lines = text.splitlines()
    blocks = []
    current_block_type = None
    current_lines = []

    def flush_block():
        if current_lines:
            content = "\n".join(current_lines)
            blocks.append({
                "type": current_block_type or "paragraph",
                "content": content
            })
            current_lines.clear()

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush_block()
            continue

        is_heading = stripped.startswith("#") or looks_like_section_heading(line)
        is_table_row = "|" in line
        is_list_item = bool(re.match(r"^([-*+]\s+|\d+[\).]\s+)", stripped))

        if is_heading:
            flush_block()
            blocks.append({
                "type": "heading",
                "content": line
            })
            current_block_type = None
        elif is_table_row:
            if current_block_type != "table":
                flush_block()
                current_block_type = "table"
            current_lines.append(line)
        elif is_list_item:
            if current_block_type != "list":
                flush_block()
                current_block_type = "list"
            current_lines.append(line)
        else:
            if current_block_type not in ("paragraph", None):
                flush_block()
            current_block_type = "paragraph"
            current_lines.append(line)

    flush_block()

    # 2. Assemble blocks into chunks
    chunks = []
    current_chunk_parts = []
    current_len = 0

    for block in blocks:
        block_content = block["content"]
        block_len = len(block_content)
        block_type = block["type"]

        # Check chunk size limit for this block
        allowed_limit = limit
        if block_type in ("table", "list"):
            allowed_limit = int(limit * CHUNK_MARGIN_FACTOR)

        # If a single block exceeds the allowed limit, split it recursively
        if block_len > allowed_limit:
            if current_chunk_parts:
                chunks.append("\n\n".join(current_chunk_parts))
                current_chunk_parts = []
                current_len = 0

            splitter = RecursiveCharacterTextSplitter(
                separators=["\n\n", "\n", ".", " ", ""],
                chunk_size=limit,
                chunk_overlap=overlap,
                length_function=len,
            )
            sub_chunks = splitter.split_text(block_content)
            chunks.extend(sub_chunks)
            continue

        # Check if adding this block exceeds the current chunk's allowed limit
        new_len = current_len + (2 if current_len > 0 else 0) + block_len
        if new_len <= allowed_limit:
            current_chunk_parts.append(block_content)
            current_len = new_len
        else:
            # Flush current chunk
            if current_chunk_parts:
                chunks.append("\n\n".join(current_chunk_parts))

            # Apply overlap: take end of previous chunk if available
            overlap_prefix = ""
            if current_chunk_parts:
                last_part = current_chunk_parts[-1]
                if len(last_part) <= overlap:
                    overlap_prefix = last_part
                else:
                    # Overlap from the end of the last part, trying to keep line boundaries
                    overlap_lines = last_part.splitlines()
                    overlap_accum = []
                    overlap_chars = 0
                    for line_part in reversed(overlap_lines):
                        if overlap_chars + len(line_part) + 1 > overlap:
                            break
                        overlap_accum.insert(0, line_part)
                        overlap_chars += len(line_part) + 1
                    overlap_prefix = "\n".join(overlap_accum)

            current_chunk_parts = []
            current_len = 0
            if overlap_prefix:
                current_chunk_parts.append(overlap_prefix)
                current_len = len(overlap_prefix)

            current_chunk_parts.append(block_content)
            current_len += (2 if current_len > 0 else 0) + block_len

    if current_chunk_parts:
        chunks.append("\n\n".join(current_chunk_parts))

    # Clean and filter chunks
    final_chunks = []
    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk:
            continue
        if re.fullmatch(r"[^a-zA-Z]+", chunk):
            continue
        # Drop tiny fragments, but NEVER discard a short chunk that carries a
        # high-value fact (a phone/email, a fee/amount, a percentage, or a named
        # role like "Principal: ..."). Such records are often < MIN_CHUNK_WORDS
        # yet are exactly what users ask for; silently dropping them removed real
        # answers from the knowledge base.
        if word_count(chunk) < MIN_CHUNK_WORDS and not is_valuable_short_chunk(chunk):
            continue
        final_chunks.append(chunk)

    return final_chunks


# Signals that a short chunk still carries a fact worth keeping (contact details,
# fees/amounts, percentages, or a named office-holder). Used to exempt such
# chunks from the MIN_CHUNK_WORDS floor so short factual records are not lost.
_VALUABLE_SHORT_PATTERNS = (
    re.compile(r"\b[\w.+-]+@[\w-]+\.[\w.-]+\b"),                 # email
    re.compile(r"(?<!\d)(?:\+?\d[\d\s().-]{7,}\d)"),             # phone number
    re.compile(r"[₹$]\s?\d", re.IGNORECASE),                     # currency amount
    re.compile(r"\b\d+(?:\.\d+)?\s?%"),                          # percentage
    re.compile(r"\b(?:rs\.?|inr|fee|fees)\b.*?\d", re.IGNORECASE),  # fee figure
    re.compile(
        r"\b(?:principal|vice[- ]?principal|hod|head of department|warden|"
        r"director|coordinator|co-ordinator|registrar|dean|secretary|"
        r"chairperson|in[- ]?charge)\b\s*[:\-]\s*\S",
        re.IGNORECASE,
    ),                                                           # named role label, e.g. "Warden - Fr. X"
)


def is_valuable_short_chunk(text: str) -> bool:
    """True when a short chunk still holds a contact, fee, percentage, or named
    role and should bypass the MIN_CHUNK_WORDS floor."""
    text = str(text or "")
    if not text.strip():
        return False
    return any(pattern.search(text) for pattern in _VALUABLE_SHORT_PATTERNS)


# Header overhead (Title/Source/Section lines) added to every embedding text.
# Reserved against the model token budget so the chunk body still fits once the
# header is prepended at embed time.
_EMBED_HEADER_TOKEN_RESERVE = 48


def split_chunks_for_embedding(chunks: list[str]) -> list[str]:
    """
    Forward-only safeguard against silent embedding truncation.

    The embedding model (BGE-base) has a 512-token cap and SentenceTransformer
    truncates silently, so a long table/list chunk loses its tail from the
    vector (measured: ~8% of chunks exceeded 512 tokens). Split any chunk whose
    tokenised length (plus the header reserve) exceeds the model budget so the
    full content is embedded across multiple chunks instead of being dropped.

    Best-effort: if the tokenizer is unavailable (e.g. model not loaded in a unit
    test), the chunks are returned unchanged — never raises.
    """
    if not chunks:
        return chunks
    try:
        from embeddings import get_embedding_model

        model = get_embedding_model()
        tokenizer = model.tokenizer
        budget = int(getattr(model, "max_seq_length", 512) or 512) - _EMBED_HEADER_TOKEN_RESERVE
    except Exception:
        return chunks
    if budget <= 0:
        return chunks

    out: list[str] = []
    for chunk in chunks:
        try:
            n_tokens = len(tokenizer.encode(chunk, add_special_tokens=True))
        except Exception:
            out.append(chunk)
            continue
        if n_tokens <= budget:
            out.append(chunk)
            continue

        approx_chars = max(200, int(len(chunk) * budget / max(n_tokens, 1)))
        splitter = RecursiveCharacterTextSplitter(
            separators=["\n\n", "\n", ". ", " ", ""],
            chunk_size=approx_chars,
            chunk_overlap=min(CHUNK_OVERLAP, approx_chars // 5),
            length_function=len,
        )
        for sub in splitter.split_text(chunk):
            sub = sub.strip()
            if not sub:
                continue
            try:
                sub_ids = tokenizer.encode(sub, add_special_tokens=True)
                if len(sub_ids) > budget:
                    # Last-resort hard cap for pathologically dense fragments.
                    sub = tokenizer.decode(
                        sub_ids[1 : budget - 1], skip_special_tokens=True
                    ).strip()
            except Exception:
                pass
            if sub:
                out.append(sub)
    return out


# =============================================================================
# HELPERS
# =============================================================================

def safe_id_text(value) -> str:
    value = str(value or "unknown").replace(" ", "_")
    return re.sub(r"[^a-zA-Z0-9_\-.]", "_", value)


def normalize_metadata_value(value, default="general"):
    if value is None:
        return default

    if isinstance(value, (str, int, float, bool)):
        return value

    return str(value)


def detect_section_title(text: str) -> str:
    text = clean_loaded_text(text)

    if not text:
        return "general"

    lines = [line.strip() for line in text.splitlines() if line.strip()]

    ignore_patterns = [
        "st. anthony",
        "handbook",
        "ever more",
        "page",
        "document:",
        # Address / institutional boilerplate that is NOT a section heading but
        # often appears at the top of a page above the real heading.
        "don bosco",
        "bomfyle",
        "shillong",
        "meghalaya",
        "affiliated",
    ]

    heading_keywords = [
        "committee",
        "cell",
        "department",
        "departments",
        "rules",
        "guidelines",
        "attendance",
        "library",
        "hostel",
        "fee",
        "fees",
        "admission",
        "examination",
        "conduct",
        "decency",
        "mission",
        "vision",
        "iqac",
        "course",
        "courses",
        "programme",
        "programmes",
        "syllabus",
        "curriculum",
        "notice",
        "news",
        "contact",
        "faculty",
        "staff",
        "quick links",
        "downloads",
        "pdf",
        # Profile / about headings (the real heading is often buried below the
        # address block, so these were previously missed in favour of boilerplate).
        "profile",
        "about",
        "overview",
        "introduction",
        "objectives",
        "history",
    ]

    def _is_contact_boilerplate(line: str) -> bool:
        low = line.lower()
        if any(tok in low for tok in (
            "phone", "fax", "tel:", "tele", "e-mail", "email", "www.",
            "http", "@", "pincode", "pin:", "po box",
        )):
            return True
        alpha = sum(c.isalpha() for c in line)
        digits = sum(c.isdigit() for c in line)
        if alpha == 0:
            return True
        # Phone numbers / dates / pincodes: digit-heavy lines are never headings.
        if digits >= 5 and digits >= alpha * 0.4:
            return True
        return False

    # 1) Strongest signal: a line near the top that carries a known section
    #    keyword and is neither boilerplate nor an address line. Scan a bit deeper
    #    (20 lines) because the real heading is often buried below the masthead.
    for line in lines[:20]:
        normalized = line.lower()
        if any(p in normalized for p in ignore_patterns) or _is_contact_boilerplate(line):
            continue
        if len(line.split()) <= 12 and any(kw in normalized for kw in heading_keywords):
            return line

    # 2) Structural heading (markdown '#', numbered, or all-caps title line).
    for line in lines[:12]:
        normalized = line.lower()
        if any(p in normalized for p in ignore_patterns) or _is_contact_boilerplate(line):
            continue
        if len(line.split()) <= 12 and looks_like_section_heading(line):
            return line

    # 3) Fallback: first short, non-boilerplate line.
    for line in lines[:8]:
        normalized = line.lower()
        if any(p in normalized for p in ignore_patterns) or _is_contact_boilerplate(line):
            continue
        if len(line.split()) <= 10:
            return line

    return " ".join(text.split()[:8])


def _table_to_text(table: list[list]) -> str:
    if not table:
        return ""

    rows = [[str(cell or "").strip() for cell in row] for row in table]
    rows = [row for row in rows if any(row)]

    if not rows:
        return ""

    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]

    headers = [cell or f"Column {i + 1}" for i, cell in enumerate(rows[0])]

    output = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]

    for row in rows[1:]:
        output.append("| " + " | ".join(row) + " |")

    return clean_loaded_text("\n".join(output))


def _extract_prose_and_tables(page) -> tuple[str, list[str], int]:
    """
    Extract a PDF page's prose and tables WITHOUT duplicating table content.

    pdfplumber's ``extract_text()`` already flattens table cells into the running
    text, so the previous code (extract_text() + extract_tables() both appended)
    embedded every table twice. Here we locate the table regions, pull the prose
    with those regions excluded, then append each table once as structured
    markdown (which retrieves better than flattened cells).

    Returns ``(prose_without_tables, [table_markdown, ...], table_count)``. On any
    failure it falls back to the plain prose + structured tables (still correct,
    just possibly with the old duplication) so a page is never dropped.
    """
    try:
        found = page.find_tables() or []
    except Exception:
        found = []

    table_texts: list[str] = []
    for tbl in found:
        try:
            rendered = _table_to_text(tbl.extract())
        except Exception:
            rendered = ""
        if rendered:
            table_texts.append(rendered)

    if not found:
        # No tables detected — plain prose, nothing to deduplicate.
        try:
            return (page.extract_text() or "", [], 0)
        except Exception:
            return ("", [], 0)

    # Build the prose with table-bbox characters removed so cell text is not
    # repeated alongside the structured tables.
    bboxes = []
    for tbl in found:
        try:
            x0, top, x1, bottom = tbl.bbox
            bboxes.append((float(x0), float(top), float(x1), float(bottom)))
        except Exception:
            continue

    def _outside_tables(obj) -> bool:
        cx = (obj.get("x0", 0) + obj.get("x1", 0)) / 2.0
        cy = (obj.get("top", 0) + obj.get("bottom", 0)) / 2.0
        for x0, top, x1, bottom in bboxes:
            if x0 <= cx <= x1 and top <= cy <= bottom:
                return False
        return True

    try:
        prose = page.filter(_outside_tables).extract_text() or ""
    except Exception:
        # Filtering failed: keep prose but drop the structured tables to avoid the
        # double-extraction (cell text already lives in the prose).
        try:
            return (page.extract_text() or "", [], 0)
        except Exception:
            return ("", table_texts, len(table_texts))

    return (prose, table_texts, len(table_texts))


def _dataframe_to_text(df: pd.DataFrame) -> str:
    if df.empty:
        return ""

    df = df.copy()

    df.columns = [
        str(col).strip()
        if str(col).strip() and not str(col).startswith("Unnamed")
        else f"Column {i + 1}"
        for i, col in enumerate(df.columns)
    ]

    rows: list[str] = []

    for row_index, row in df.iterrows():
        parts = []

        for col in df.columns:
            value = row[col]

            if pd.notna(value) and str(value).strip():
                parts.append(f"{col}: {str(value).strip()}")

        if parts:
            rows.append(f"Row {row_index + 1}: " + " | ".join(parts))

    return clean_loaded_text("\n".join(rows))


def preprocess_image_for_ocr(image):
    image = image.convert("RGB")
    image = ImageOps.grayscale(image)
    image = image.filter(ImageFilter.SHARPEN)

    return image


def run_ocr_on_image(image) -> str:
    if Image is None or ImageOps is None or ImageFilter is None or pytesseract is None:
        print("OCR skipped: Pillow and/or pytesseract is not available.")
        return ""

    try:
        image = preprocess_image_for_ocr(image)
        text = pytesseract.image_to_string(image)

        return clean_loaded_text(text)

    except Exception as e:
        print(f"OCR failed: {e}")
        return ""


# =============================================================================
# FILE LOADERS
# =============================================================================

def _boilerplate_key(line: str) -> str:
    """Normalize a line for repeated-boilerplate detection (digits -> #)."""
    return re.sub(r"\s+", " ", re.sub(r"\d+", "#", str(line or "").strip().lower())).strip()


def strip_repeated_pdf_boilerplate(documents: list[Document]) -> list[Document]:
    """Remove running headers/footers/page-numbers that repeat across PDF pages.

    Conservative by design — a line is treated as boilerplate only when it is
    short (<= 12 words) AND its normalized form (page numbers folded to '#')
    appears on at least 3 pages AND on >= 60% of pages. Real content almost never
    repeats verbatim on most pages, whereas running headers/footers always do.
    This keeps repeated institutional headers/footers out of every chunk, which
    is a common cross-topic collision source.
    """
    if len(documents) < 4:  # too few pages to judge repetition safely
        return documents

    n_pages = len(documents)
    page_keys: list[list[tuple[str, str]]] = []  # per page: list of (raw_line, key)
    line_pages: dict[str, set[int]] = {}

    for idx, doc in enumerate(documents):
        seen_on_page: set[str] = set()
        rows: list[tuple[str, str]] = []
        for raw in str(doc.page_content or "").splitlines():
            key = _boilerplate_key(raw)
            rows.append((raw, key))
            if not key or len(raw.split()) > 12:
                continue
            if key in seen_on_page:
                continue
            seen_on_page.add(key)
            line_pages.setdefault(key, set()).add(idx)
        page_keys.append(rows)

    threshold = max(3, int(0.6 * n_pages))
    boilerplate = {key for key, pages in line_pages.items() if len(pages) >= threshold}
    if not boilerplate:
        return documents

    for doc, rows in zip(documents, page_keys):
        kept = [raw for raw, key in rows if not (key and key in boilerplate)]
        doc.page_content = clean_loaded_text("\n".join(kept))

    return documents


def load_pdf_bytes(
    file_bytes: bytes,
    filename: str,
    job_id: str | None = None,
    extra_metadata: dict | None = None,
) -> list[Document]:
    documents: list[Document] = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        total_pages = len(pdf.pages)
        # Memory-safety cap: a crafted PDF can claim a huge page count and exhaust
        # memory/CPU during per-page rendering+OCR. Process at most MAX_PDF_PAGES
        # (override via env); the rest are skipped with a clear warning.
        pages_to_process = pdf.pages
        if total_pages > MAX_PDF_PAGES:
            print(
                f"[PDF] {filename}: {total_pages} pages exceeds MAX_PDF_PAGES="
                f"{MAX_PDF_PAGES}; processing the first {MAX_PDF_PAGES} pages only."
            )
            pages_to_process = pdf.pages[:MAX_PDF_PAGES]
        pdf_metadata = pdf.metadata or {}
        pdf_metadata = {str(k): str(v) for k, v in pdf_metadata.items() if v is not None}
        if extra_metadata:
            pdf_metadata.update({str(k): str(v) for k, v in extra_metadata.items() if v is not None})
        pdf_metadata_text = " ".join(str(value or "") for value in pdf_metadata.values())
        pdf_document_date = metadata_date_value(pdf_metadata, pdf_metadata_text)

        for page_index, page in enumerate(pages_to_process, start=1):
            if job_id:
                wait_if_paused(job_id)
                check_crawl_cancelled(job_id)

            page_parts: list[str] = []
            tables_extracted = 0
            ocr_used = False

            prose, tables, tables_extracted = _extract_prose_and_tables(page)

            if prose:
                prose = clean_loaded_text(prose)
                if prose:
                    page_parts.append(prose)

            for table_text in tables:
                table_text = clean_loaded_text(table_text)
                if table_text:
                    page_parts.append(table_text)

            combined_before_ocr = clean_loaded_text("\n\n".join(page_parts))

            if len(combined_before_ocr.split()) < PDF_OCR_FALLBACK_MIN_WORDS:
                try:
                    if job_id:
                        wait_if_paused(job_id)
                        check_crawl_cancelled(job_id)
                    page_image = page.to_image(resolution=200).original
                    ocr_text = run_ocr_on_image(page_image)

                    if ocr_text:
                        page_parts.append(ocr_text)
                        ocr_used = True

                except Exception as e:
                    print(f"PDF OCR fallback failed on {filename}, page {page_index}: {e}")

            combined = clean_loaded_text("\n\n".join(page_parts))

            if not combined:
                continue

            section_title = detect_section_title(combined)
            document_year = metadata_year_value(
                filename,
                pdf_document_date,
                pdf_metadata_text,
                section_title,
                combined[:3000],
            )

            documents.append(
                Document(
                    page_content=combined,
                    metadata={
                        "filename": filename,
                        "source_filename": filename,
                        "source_url": "",
                        "title": filename,
                        "page": page_index,
                        "total_pages": total_pages,
                        "file_type": "pdf",
                        "source_type": "pdf",
                        "document_year": document_year,
                        "document_date": pdf_document_date,
                        "crawl_timestamp": utc_timestamp(),
                        "section_title": section_title,
                        "tables_extracted": tables_extracted,
                        "ocr_used": ocr_used,
                        "text_chars": len(combined),
                        "is_toc": detect_toc(combined),
                    },
                )
            )

    # Strip running headers/footers/page-numbers that repeat across pages so the
    # same institutional boilerplate does not enter every chunk.
    documents = strip_repeated_pdf_boilerplate(documents)
    return documents


def load_image_bytes(file_bytes: bytes, filename: str, url: str = None) -> list[Document]:
    if Image is None:
        print("Image OCR skipped: Pillow is not available.")
        return []

    check_name = filename.lower()
    check_url = (url or "").lower()
    # Only skip clear non-content decorations (logos/icons/avatars). Person/path
    # terms like "profile", "staff", "user" were previously dropped too, which
    # silently discarded legitimate notices/posters served from such paths. True
    # headshots with no text simply OCR to empty and are dropped below anyway, so
    # this narrow list avoids losing real content. Small logos/icons are still
    # caught by the size floor below.
    skip_decoration_terms = ["logo", "icon", "avatar"]

    if any(term in check_name for term in skip_decoration_terms) or \
       any(term in check_url for term in skip_decoration_terms):
        print(f"[IMAGE] Skipped logo/icon/avatar image: {filename}")
        return []

    try:
        image = Image.open(io.BytesIO(file_bytes)).convert("RGB")
    except Exception as e:
        print(f"[IMAGE] Failed to open image {filename}: {e}")
        return []

    width, height = image.size
    area = width * height

    if width < 400 or height < 300 or area < 120000:
        print(f"[IMAGE] Skipped small image ({width}x{height}): {filename}")
        return []

    # Note: square images are NOT skipped here. Notices/posters/infographics are
    # often square, and dropping them by aspect ratio discarded real text. Images
    # with no text simply OCR to empty and are dropped below.
    text = run_ocr_on_image(image)

    if not text:
        return []

    return [
        Document(
            page_content=text,
            metadata={
                "filename": filename,
                "page": 1,
                "total_pages": 1,
                "file_type": "image_ocr",
                "section_title": detect_section_title(text),
            },
        )
    ]


def validate_zip_archive_safety(file_path: str) -> None:
    """
    Guard ZIP-based Office files (DOCX/XLSX/PPTX) against decompression bombs by
    inspecting the central directory BEFORE any member is decompressed. Raises
    ValueError with a clear message when the archive exceeds the member count,
    per-member size, total uncompressed size, or compression-ratio limits.

    Reads only metadata (sizes from the ZIP directory), so it is cheap and never
    expands the bomb. A non-ZIP file is ignored here — the format-specific loader
    raises its own error.
    """
    import zipfile

    name = os.path.basename(file_path)
    try:
        archive = zipfile.ZipFile(file_path)
    except (zipfile.BadZipFile, OSError):
        return  # not a valid zip; let the real loader surface the error

    with archive:
        infos = archive.infolist()
        if len(infos) > MAX_ARCHIVE_MEMBERS:
            raise ValueError(
                f"Refusing {name}: archive has {len(infos)} members, exceeding the "
                f"{MAX_ARCHIVE_MEMBERS}-member limit (possible zip bomb)."
            )

        total_uncompressed = 0
        total_compressed = 0
        for info in infos:
            if info.file_size > MAX_ARCHIVE_MEMBER_BYTES:
                raise ValueError(
                    f"Refusing {name}: member {info.filename!r} expands to "
                    f"{info.file_size} bytes, exceeding the per-member "
                    f"{MAX_ARCHIVE_MEMBER_BYTES}-byte limit (possible zip bomb)."
                )
            total_uncompressed += info.file_size
            total_compressed += info.compress_size
            if total_uncompressed > MAX_ARCHIVE_TOTAL_BYTES:
                raise ValueError(
                    f"Refusing {name}: total uncompressed size exceeds the "
                    f"{MAX_ARCHIVE_TOTAL_BYTES}-byte limit (possible zip bomb)."
                )

        if total_compressed > 0 and total_uncompressed > _ARCHIVE_RATIO_MIN_BYTES:
            ratio = total_uncompressed / total_compressed
            if ratio > MAX_ARCHIVE_COMPRESSION_RATIO:
                raise ValueError(
                    f"Refusing {name}: compression ratio {ratio:.0f}x exceeds the "
                    f"{MAX_ARCHIVE_COMPRESSION_RATIO:.0f}x limit (possible zip bomb)."
                )


def load_docx_file(file_path: str) -> list[Document]:
    validate_zip_archive_safety(file_path)
    import docx
    doc = docx.Document(file_path)
    text = clean_loaded_text("\n".join([p.text for p in doc.paragraphs if p.text.strip()]))
    
    if not text:
        return []

    return [
        Document(
            page_content=text,
            metadata={
                "page": 1,
                "total_pages": 1,
                "file_type": "docx",
                "section_title": detect_section_title(text),
            },
        )
    ]


def load_txt_file(file_path: str) -> list[Document]:
    documents = TextLoader(file_path, encoding="utf-8").load()

    for i, doc in enumerate(documents, start=1):
        doc.page_content = clean_loaded_text(doc.page_content)
        doc.metadata["page"] = doc.metadata.get("page", i)
        doc.metadata["total_pages"] = doc.metadata.get("total_pages", len(documents))
        doc.metadata["file_type"] = "txt"
        doc.metadata["section_title"] = detect_section_title(doc.page_content)

    return [doc for doc in documents if doc.page_content]


def load_html_file(file_path: str) -> list[Document]:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        html = f.read()
        
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["nav", "footer", "header", "aside", "script", "style"]):
        tag.decompose()
        
    text = clean_loaded_text(soup.get_text(separator="\n", strip=True))
    
    if not text:
        return []

    return [
        Document(
            page_content=text,
            metadata={
                "page": 1,
                "total_pages": 1,
                "file_type": "html",
                "section_title": detect_section_title(text),
            },
        )
    ]


def load_csv_file(file_path: str) -> list[Document]:
    df = pd.read_csv(file_path)
    text = _dataframe_to_text(df)

    if not text:
        return []

    return [
        Document(
            page_content=text,
            metadata={
                "source": file_path,
                "filename": os.path.basename(file_path),
                "page": 1,
                "total_pages": 1,
                "file_type": "csv",
                "section_title": detect_section_title(text),
            },
        )
    ]


def load_markdown_file(file_path: str) -> list[Document]:
    documents = TextLoader(file_path, encoding="utf-8").load()

    for i, doc in enumerate(documents, start=1):
        doc.page_content = clean_loaded_text(doc.page_content)
        doc.metadata["page"] = doc.metadata.get("page", i)
        doc.metadata["total_pages"] = doc.metadata.get("total_pages", len(documents))
        doc.metadata["file_type"] = "markdown"
        doc.metadata["section_title"] = detect_section_title(doc.page_content)

    return [doc for doc in documents if doc.page_content]


def load_json_file(file_path: str) -> list[Document]:
    with open(file_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    text = clean_loaded_text(json.dumps(data, indent=2, ensure_ascii=False))

    if not text:
        return []

    return [
        Document(
            page_content=text,
            metadata={
                "source": file_path,
                "filename": os.path.basename(file_path),
                "page": 1,
                "total_pages": 1,
                "file_type": "json",
                "section_title": detect_section_title(text),
            },
        )
    ]


def load_xlsx_file(file_path: str) -> list[Document]:
    validate_zip_archive_safety(file_path)
    documents: list[Document] = []
    excel_file = pd.ExcelFile(file_path)
    total_sheets = len(excel_file.sheet_names)

    for i, sheet_name in enumerate(excel_file.sheet_names, start=1):
        df = pd.read_excel(file_path, sheet_name=sheet_name)
        text = _dataframe_to_text(df)

        if not text:
            continue

        documents.append(
            Document(
                page_content=text,
                metadata={
                    "source": file_path,
                    "filename": os.path.basename(file_path),
                    "page": i,
                    "total_pages": total_sheets,
                    "file_type": "xlsx",
                    "sheet_name": sheet_name,
                    "section_title": str(sheet_name),
                },
            )
        )

    return documents


def load_db_dump_file(file_path: str) -> list[Document]:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = clean_loaded_text(f.read())

    if not text:
        return []

    return [
        Document(
            page_content=text,
            metadata={
                "source": file_path,
                "filename": os.path.basename(file_path),
                "page": 1,
                "total_pages": 1,
                "file_type": "db_dump",
                "section_title": detect_section_title(text),
            },
        )
    ]


# =============================================================================
# WEBSITE CRAWLING + WEBSITE PDF INGESTION
# =============================================================================

def normalize_url(url: str, keep_query: bool = False) -> str:
    parsed = urlparse(str(url or "").strip())
    clean = parsed._replace(fragment="") if keep_query else parsed._replace(fragment="", query="")
    return clean.geturl().rstrip("/")


def get_base_domain(url: str) -> str:
    parsed = urlparse(url)
    return parsed.netloc.lower().replace("www.", "")


def is_private_ip(url: str) -> bool:
    parsed = urlparse(url)
    hostname = (parsed.hostname or "").lower()
    if hostname in ("localhost", "127.0.0.1", "0.0.0.0", "::1"):
        return True
    # IPv4 loopback (127/8) and link-local (169.254/16, incl. cloud metadata).
    if hostname.startswith(("127.", "169.254.")):
        return True
    # IPv4 RFC1918 private ranges.
    if re.match(r"^(10\.|172\.(1[6-9]|2[0-9]|3[01])\.|192\.168\.)", hostname):
        return True
    # IPv6 loopback / unique-local (fc00::/7) / link-local (fe80::/10). Gated on
    # ":" so ordinary hostnames like "fcollege.com" are not misclassified.
    if ":" in hostname and hostname.startswith(("::1", "fc", "fd", "fe80")):
        return True
    return False


def is_same_domain(url: str, base_domain: str) -> bool:
    parsed = urlparse(url)
    netloc = parsed.netloc.lower().replace("www.", "")

    return netloc == base_domain


def is_pdf_url(url: str) -> bool:
    lower = str(url or "").lower()
    parsed = urlparse(lower)
    return parsed.path.endswith(".pdf") or ".pdf" in lower or "pdf=" in lower


def is_document_url(url: str) -> bool:
    lower = str(url or "").lower()
    path = urlparse(lower).path
    if any(path.endswith(ext) for ext in DOCUMENT_EXTENSIONS):
        return True
    if any(ext in lower for ext in DOCUMENT_EXTENSIONS):
        return True
    document_markers = (
        "download.php",
        "file=",
        "doc=",
        "pdf=",
        "attachment",
        "uploads",
        "resources",
    )
    return any(marker in lower for marker in document_markers)


def is_supported_document_response(url: str, content_type: str) -> bool:
    lower_url = str(url or "").lower()
    lower_type = str(content_type or "").lower()

    if "application/pdf" in lower_type or is_pdf_url(lower_url):
        return True

    parsed = urlparse(lower_url)
    ext = os.path.splitext(parsed.path)[1]
    if ext in DOCUMENT_EXTENSIONS:
        return True

    supported_types = {
        "text/plain",
        "text/csv",
        "text/markdown",
        "text/html",
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/vnd.ms-excel",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/json",
        "application/sql",
    }
    for st in supported_types:
        if st in lower_type:
            return True

    return False


# Hard domain allow-list (anti-collision + anti-SSRF). EduBot is a single-college
# assistant, so by default the crawler may ONLY visit the official college domain.
# Override via env ALLOWED_CRAWL_DOMAINS (comma-separated registrable domains).
# Setting it empty disables the hard-lock (generic behaviour for other deployments).
ALLOWED_CRAWL_DOMAINS = {
    d.strip().lower().removeprefix("www.")
    for d in os.getenv("ALLOWED_CRAWL_DOMAINS", "anthonys.ac.in").split(",")
    if d.strip()
}


def _url_host(url: str) -> str:
    host = (urlparse(str(url or "")).hostname or "").lower()
    return host.removeprefix("www.")


def is_domain_allowed(url: str) -> bool:
    """True when the URL's host is within the configured college allow-list.

    Subdomains of an allowed domain are permitted (e.g. admissions.anthonys.ac.in).
    When the allow-list is empty the hard-lock is disabled and all hosts pass.
    """
    if not ALLOWED_CRAWL_DOMAINS:
        return True
    host = _url_host(url)
    if not host:
        return False
    return any(host == d or host.endswith("." + d) for d in ALLOWED_CRAWL_DOMAINS)


def is_allowed_crawl_url(url: str) -> bool:
    if not url:
        return False

    lower = str(url).lower().strip()

    if is_private_ip(lower):
        return False

    if lower.startswith(("mailto:", "tel:", "javascript:", "#")):
        return False

    if any(lower.endswith(ext) for ext in WEBSITE_BLOCKED_EXTENSIONS):
        return False

    if not lower.startswith(("http://", "https://")):
        return False

    # Hard domain lock: every crawled/queued/linked URL must be inside the college
    # allow-list. Enforced here so it cannot be bypassed by the same_domain_only
    # flag and applies to both crawler backends.
    return is_domain_allowed(lower)


def is_html_response(content_type: str, body: str = "") -> bool:
    lower_type = str(content_type or "").lower()
    if "html" in lower_type or "xml" in lower_type:
        return True

    if lower_type and not lower_type.startswith("text/"):
        return False

    sample = str(body or "")[:500].lower()
    return "<html" in sample or "<!doctype html" in sample


def safe_pdf_filename_from_url(url: str, index: int = 1) -> str:
    parsed = urlparse(url)
    filename = os.path.basename(parsed.path.rstrip("/"))

    query = parse_qs(parsed.query)
    for key in ("file", "filename", "doc", "pdf", "attachment"):
        value = query.get(key, [""])[0]
        if value:
            filename = os.path.basename(unquote(value))
            break

    if not filename:
        filename = f"website_pdf_{index}.pdf"

    filename = re.sub(r"[^a-zA-Z0-9_.-]", "_", filename)

    img_extensions = {".png", ".jpg", ".jpeg", ".webp"}
    has_img_ext = any(filename.lower().endswith(ext) for ext in img_extensions)
    
    if not filename.lower().endswith(".pdf") and not has_img_ext:
        filename += ".pdf"

    return filename


def safe_document_filename_from_url(url: str, content_type: str = "", index: int = 1) -> str:
    parsed = urlparse(url)
    filename = os.path.basename(parsed.path.rstrip("/"))

    query = parse_qs(parsed.query)
    for key in ("file", "filename", "doc", "attachment"):
        value = query.get(key, [""])[0]
        if value:
            filename = os.path.basename(unquote(value))
            break

    import mimetypes
    ext = os.path.splitext(filename)[1].lower()
    if not ext:
        ext = mimetypes.guess_extension(content_type.split(";")[0]) or ".doc"
        filename += ext

    filename = re.sub(r"[^a-zA-Z0-9_.-]", "_", filename)
    if not filename or filename.startswith((".", "_")):
        filename = f"website_doc_{index}{ext or '.doc'}"

    return filename


def readable_website_source_name(url: str, title: str = "") -> str:
    title = clean_loaded_text(title)
    if title and len(title) <= 120:
        return title

    parsed = urlparse(str(url or ""))
    path = parsed.path.strip("/")
    if path:
        return path

    return parsed.netloc or str(url or "website")


def dedupe_loaded_lines(text: str) -> str:
    lines = []
    seen = set()

    for line in str(text or "").splitlines():
        cleaned = re.sub(r"\s+", " ", line).strip()

        if not cleaned:
            if lines and lines[-1]:
                lines.append("")
            continue

        key = cleaned.lower()
        if key in seen:
            continue

        seen.add(key)
        lines.append(cleaned)

    while lines and not lines[-1]:
        lines.pop()

    return "\n".join(lines)


def is_hidden_html_element(tag) -> bool:
    if not getattr(tag, "attrs", None):
        return False

    if tag.has_attr("hidden"):
        return True

    if str(tag.get("aria-hidden", "")).lower() == "true":
        return True

    style = str(tag.get("style", "")).replace(" ", "").lower()
    return "display:none" in style or "visibility:hidden" in style


def extract_visible_text_from_website_html(html: str, url: str) -> str:
    soup = BeautifulSoup(html or "", "lxml")

    for tag in soup(["script", "style", "noscript", "svg", "iframe", "canvas"]):
        tag.decompose()

    for tag in soup.find_all(is_hidden_html_element):
        tag.decompose()

    lines = []
    title = soup.title.get_text(" ", strip=True) if soup.title else ""
    if title:
        lines.append(f"Page title: {title}")

    content_tags = (
        "h1", "h2", "h3", "h4", "h5", "h6",
        "p", "li", "td", "th", "caption",
        "article", "section", "main",
        "a", "button",
    )

    for tag in soup.find_all(content_tags):
        text = re.sub(r"\s+", " ", tag.get_text(" ", strip=True)).strip()
        if not text:
            continue

        if tag.name == "a":
            href = tag.get("href", "")
            absolute_url = normalize_url(
                urljoin(url, href),
                keep_query=is_document_url(urljoin(url, href)),
            )
            if is_allowed_crawl_url(absolute_url):
                lines.append(f"- {text}: {absolute_url}")
            else:
                lines.append(f"- {text}")
            continue

        if tag.name == "button":
            lines.append(f"Button: {text}")
            continue

        if tag.name == "li":
            lines.append(f"- {text}")
            continue

        lines.append(text)

    return clean_loaded_text(dedupe_loaded_lines("\n".join(lines)))


def extract_clean_text_from_website_html(html: str, url: str) -> str:
    extracted = trafilatura.extract(
        html or "",
        url=url,
        include_tables=True,
        include_links=False,
        include_comments=False,
        no_fallback=False,
    )

    text_parts = []
    if extracted:
        text_parts.append(extracted)

    soup = BeautifulSoup(html or "", "lxml")

    for tag in soup(["script", "style", "aside", "form", "iframe", "noscript", "svg"]):
        tag.decompose()

    for tag in soup.find_all(is_hidden_html_element):
        tag.decompose()

    soup_text = soup.get_text(separator="\n")
    if soup_text:
        text_parts.append(soup_text)

    visible_text = extract_visible_text_from_website_html(html, url)
    if visible_text:
        text_parts.append(visible_text)

    return clean_loaded_text(dedupe_loaded_lines("\n\n".join(text_parts)))


def extract_visible_links_text_from_html(html: str, current_url: str, base_domain: str) -> str:
    soup = BeautifulSoup(html or "", "html.parser")
    
    title = soup.title.string.strip() if soup.title else "No Title"
    output_lines = [f"Page title: {title}", f"Source page: {current_url}", ""]
    output_lines.append("Website links / quick links found on this page:")
    
    seen_links = set()
    link_sections = []
    relevant_headers = [
        "quick links", "useful links", "important links", "student links", 
        "downloads", "notices", "resources", "navigation", "menu"
    ]

    # First, look for links within sections that have relevant headers (Task 6)
    for header_tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "strong", "b", "span"]):
        header_text = header_tag.get_text().strip().lower()
        if any(keyword in header_text for keyword in relevant_headers):
            parent = header_tag.parent
            section_links = []
            # Look for <a> tags in siblings or nearby parents
            for link in parent.find_all("a", href=True):
                label = link.get_text(" ", strip=True)
                raw_href = urljoin(current_url, link["href"])
                href = normalize_url(raw_href, keep_query=is_document_url(raw_href))
                
                if not label or not is_allowed_crawl_url(href) or not is_same_domain(href, base_domain):
                    continue
                
                key = f"{label}|{href}"
                if key not in seen_links:
                    seen_links.add(key)
                    prefix = " [PDF]" if is_pdf_url(href) else ""
                    section_links.append(f"- {label}{prefix}: {href}")
            
            if section_links:
                link_sections.append(f"\n### {header_tag.get_text().strip()}\n" + "\n".join(section_links))

    # Then find all other links not already captured (Task 2 & 4)
    other_links = []
    for tag in soup.find_all("a", href=True):
        label = tag.get_text(" ", strip=True)
        raw_href = urljoin(current_url, tag["href"])
        href = normalize_url(raw_href, keep_query=is_document_url(raw_href))

        if not label or not is_allowed_crawl_url(href) or not is_same_domain(href, base_domain):
            continue

        key = f"{label}|{href}"
        if key not in seen_links:
            seen_links.add(key)
            prefix = " [PDF]" if is_pdf_url(href) else ""
            other_links.append(f"- {label}{prefix}: {href}")
            
    if link_sections:
        output_lines.extend(link_sections)
    
    if other_links:
        if link_sections:
            output_lines.append("\nOther navigation links:")
        output_lines.extend(other_links)

    if len(output_lines) <= 4: # No actual links found beyond headers
        return ""

    return "\n".join(output_lines)


def build_pdf_link_document(
    pdf_url: str,
    pdf_title: str,
    found_on_url: str,
    start_url: str,
    pdf_index: int,
) -> Document:
    parsed = urlparse(pdf_url)
    ext = os.path.splitext(parsed.path)[1].lower()
    is_img = ext in {".png", ".jpg", ".jpeg", ".webp"}
    doc_type_label = "Image document/notice" if is_img else "PDF document"
    filename = os.path.basename(parsed.path) or f"website_doc_{pdf_index}{ext}"

    text = clean_loaded_text(
        f"""
{doc_type_label} found on the college website.

Title: {pdf_title}
Filename: {filename}
URL: {pdf_url}
Found on webpage: {found_on_url}
Website source: {start_url}
"""
    )

    return Document(
        page_content=text,
        metadata={
            "filename": pdf_title or filename,
            "source_filename": pdf_title or filename,
            "source_url": pdf_url,
            "found_on_url": found_on_url,
            "source_pdf_filename": filename if not is_img else "",
            "crawl_base_url": start_url,
            "title": pdf_title,
            "pdf_title": pdf_title,
            "page": 1,
            "total_pages": 1,
            "file_type": "website_image" if is_img else "website_document",
            "source_type": "website_image" if is_img else "website_document",
            "section_title": "Website Images" if is_img else "Website PDF Documents",
            "is_toc": False,
            "document_year": 2026,
            "document_date": "2026-06-12",
            "scope": "official",
            "status": "active",
            "deleted": False,
        },
    )


def load_website(
    start_url: str,
    max_pages: int = 50,
    delay_seconds: float = 1.5,
    include_pdfs: bool = True,
    max_pdfs: int = 200,
    same_domain_only: bool = True,
    crawl_documents: bool = True,
    max_depth: int = 3,
    job_id: str | None = None,
) -> list[Document]:
    # Try Crawl4AI first
    try:
        from crawl4ai_crawler import crawl_with_crawl4ai, CRAWL4AI_AVAILABLE, CRAWL4AI_IMPORT_ERROR
        if not CRAWL4AI_AVAILABLE:
            reason = CRAWL4AI_IMPORT_ERROR or "package not importable"
            raise ImportError(
                f"Crawl4AI is not available in the active interpreter ({sys.executable}): {reason}. "
                "Ensure the backend runs inside backend/.venv and that "
                "'pip install crawl4ai && python -m playwright install chromium' has been run."
            )
        
        print(f"INFO  Crawl started: {start_url}")
        import asyncio
        documents = asyncio.run(
            crawl_with_crawl4ai(
                start_url=start_url,
                max_pages=max_pages,
                delay_seconds=delay_seconds,
                include_pdfs=include_pdfs,
                max_pdfs=max_pdfs,
                same_domain_only=same_domain_only,
                crawl_documents=crawl_documents,
                max_depth=max_depth,
                job_id=job_id
            )
        )
        print("INFO  Crawl completed")
        return documents
    except Exception as e:
        print(f"WARNING  Crawl4AI failed: {e}. Falling back to legacy crawler.")
        if job_id:
            update_crawl_job(
                job_id,
                errors=f"Crawl4AI failed: {e}. Falling back to legacy crawler."
            )
        return load_website_legacy(
            start_url=start_url,
            max_pages=max_pages,
            delay_seconds=delay_seconds,
            include_pdfs=include_pdfs,
            max_pdfs=max_pdfs,
            same_domain_only=same_domain_only,
            crawl_documents=crawl_documents,
            max_depth=max_depth,
            job_id=job_id
        )


def load_website_legacy(
    start_url: str,
    max_pages: int = 50,
    delay_seconds: float = 1.5,
    include_pdfs: bool = True,
    max_pdfs: int = 200,
    same_domain_only: bool = True,
    crawl_documents: bool = True,
    max_depth: int = 3,
    job_id: str | None = None,
) -> list[Document]:
    start_url = normalize_url(start_url)
    parsed = urlparse(start_url)

    if not parsed.scheme.startswith("http") or not parsed.netloc:
        raise ValueError("Invalid website URL.")

    base_domain = get_base_domain(start_url)

    visited: set[str] = set()
    queued: set[str] = {start_url}
    queue: list[tuple[str, int]] = [(start_url, 0)]

    documents: list[Document] = []
    failed_pages: list[dict] = []
    pdf_found_on_urls: dict[str, str] = {}

    total_links_extracted = 0
    pdfs_loaded = 0
    docs_downloaded = 0
    pdf_links_recorded: set[str] = set()

    session = create_robust_session(CRAWLER_HEADERS)

    rp = RobotFileParser()
    rp.set_url(urljoin(start_url, "/robots.txt"))
    robots_ready = False
    try:
        rp.read()
        robots_ready = True
    except Exception:
        pass

    print("[WEBSITE CRAWL START]")
    print("Start URL :", start_url)
    print("Domain    :", base_domain)

    if job_id:
        update_crawl_job(
            job_id,
            url=start_url,
            status="crawling",
            pages_found=1,
            current_stage="queued"
        )

    while queue and len(visited) < max_pages:
        url_tuple = queue.pop(0)
        url = normalize_url(url_tuple[0])
        depth = url_tuple[1]
        queued.discard(url)

        if url in visited:
            continue

        if not is_allowed_crawl_url(url):
            continue

        from crawler import is_excluded_url
        excluded, matched_pattern = is_excluded_url(url)
        if excluded:
            print(f"[Crawler] Skipped excluded URL: {url}")
            print(f'Reason: Matched exclusion pattern "{matched_pattern}"')
            continue

        if robots_ready and not rp.can_fetch(CRAWLER_HEADERS.get("User-Agent"), url):
            print(f"[WEBSITE] Blocked by robots.txt: {url}")
            continue

        if depth > max_depth:
            continue

        if same_domain_only and not is_same_domain(url, base_domain):
            continue

        is_pdf_or_doc = is_document_url(url) or is_pdf_url(url)

        # Checkpoint: before fetching a page or downloading a doc
        try:
            wait_if_paused(job_id)
            check_crawl_cancelled(job_id)
        except ValueError as e:
            print(f"[WEBSITE] Cancelled: {e}")
            break

        if is_pdf_or_doc:
            if should_skip_current_document(job_id):
                print(f"[WEBSITE] Skipping document before download: {url}")
                if job_id:
                    update_crawl_job(job_id, skipped_urls=url)
                continue
            if job_id:
                update_crawl_job(
                    job_id,
                    current_url=url,
                    current_type="pdf" if is_pdf_url(url) else "document",
                    current_stage="fetching"
                )
        else:
            if should_skip_current_page(job_id):
                print(f"[WEBSITE] Skipping page before fetch: {url}")
                if job_id:
                    update_crawl_job(job_id, skipped_urls=url)
                continue
            if job_id:
                update_crawl_job(
                    job_id,
                    current_url=url,
                    current_type="page",
                    current_stage="fetching"
                )

        visited.add(url)

        print(f"[WEBSITE] Visiting ({len(visited)}/{max_pages}): {url}")

        try:
            response = fetch_with_ssl_fallback(session, url, timeout=REQUEST_TIMEOUT_SECONDS)
            content_type = response.headers.get("content-type", "").lower()
            last_modified = response.headers.get("last-modified") or response.headers.get("Last-Modified")
            extra_meta = {}
            if last_modified:
                extra_meta["last_modified"] = last_modified

            if response.status_code != 200:
                err_msg = f"Failed to fetch {url} (HTTP {response.status_code})"
                if job_id:
                    update_crawl_job(job_id, errors=err_msg)
                failed_pages.append(
                    {
                        "url": url,
                        "status": response.status_code,
                    }
                )
                continue

            # Checkpoint: after fetching a page or downloading a doc
            try:
                wait_if_paused(job_id)
                check_crawl_cancelled(job_id)
            except ValueError as e:
                print(f"[WEBSITE] Cancelled: {e}")
                break

            if is_pdf_or_doc:
                if should_skip_current_document(job_id):
                    print(f"[WEBSITE] Skipping document after download: {url}")
                    if job_id:
                        update_crawl_job(job_id, skipped_urls=url)
                    continue
                if job_id:
                    update_crawl_job(job_id, current_stage="extracting")
            else:
                if should_skip_current_page(job_id):
                    print(f"[WEBSITE] Skipping page after fetch: {url}")
                    if job_id:
                        update_crawl_job(job_id, skipped_urls=url)
                    continue
                if job_id:
                    update_crawl_job(job_id, current_stage="extracting")

            is_doc = is_document_url(url)
            is_pdf_response = "application/pdf" in content_type or is_pdf_url(url)

            # -----------------------------------------------------------------
            # DOCUMENT CONTENT INGESTION (PDF, DOCX, etc)
            # -----------------------------------------------------------------
            if is_pdf_response or is_doc:
                if not crawl_documents:
                    continue

                if pdfs_loaded >= max_pdfs:
                    print(f"[WEBSITE DOC] Skipped because limit reached: {url}")
                    continue

                try:
                    # Checkpoint: before downloading/processing the doc
                    try:
                        wait_if_paused(job_id)
                        check_crawl_cancelled(job_id)
                    except ValueError as e:
                        print(f"[WEBSITE] Cancelled: {e}")
                        break

                    if should_skip_current_document(job_id):
                        if job_id:
                            update_crawl_job(job_id, skipped_urls=url)
                        continue

                    pdfs_loaded += 1
                    pdf_filename = safe_pdf_filename_from_url(url, pdfs_loaded)
                    
                    local_path = UPLOAD_DIR / pdf_filename
                    
                    # Download and save if not exists
                    if not local_path.exists():
                        with open(local_path, "wb") as f:
                            f.write(response.content)
                        docs_downloaded += 1

                    # Checkpoint: after downloading / before ingesting
                    try:
                        wait_if_paused(job_id)
                        check_crawl_cancelled(job_id)
                    except ValueError as e:
                        print(f"[WEBSITE] Cancelled: {e}")
                        break

                    if should_skip_current_document(job_id):
                        if job_id:
                            update_crawl_job(job_id, skipped_urls=url)
                        continue

                    if job_id:
                        update_crawl_job(job_id, current_stage="ingesting")

                    if not is_pdf_response:
                        # Non-PDF documents need to be loaded from path
                        pdf_docs = load_file_from_path(str(local_path), job_id=job_id, url=url, extra_metadata=extra_meta)
                    else:
                        pdf_docs = load_pdf_bytes(
                            file_bytes=response.content,
                            filename=pdf_filename,
                            job_id=job_id,
                            extra_metadata=extra_meta,
                        )

                    for pdf_doc in pdf_docs:
                        found_on_url = pdf_found_on_urls.get(url, "")
                        doc_title = pdf_filename
                        pdf_doc.metadata["filename"] = pdf_filename
                        pdf_doc.metadata["source_filename"] = pdf_filename
                        pdf_doc.metadata["source_url"] = url
                        pdf_doc.metadata["title"] = doc_title
                        pdf_doc.metadata["found_on_url"] = found_on_url
                        pdf_doc.metadata["source_pdf_filename"] = pdf_filename
                        pdf_doc.metadata["crawl_base_url"] = start_url
                        pdf_doc.metadata["file_type"] = "website_pdf"
                        pdf_doc.metadata["source_type"] = "pdf"
                        pdf_doc.metadata["pdf_title"] = pdf_filename
                        doc_year = pdf_doc.metadata.get("document_year")
                        doc_date = pdf_doc.metadata.get("document_date")
                        if not doc_year or doc_year == "general":
                            doc_year = metadata_year_value(
                                doc_date,
                                pdf_filename,
                                url,
                                pdf_doc.page_content[:3000],
                            )
                        if doc_year and doc_year != "general":
                            try:
                                doc_year = int(doc_year)
                            except ValueError:
                                doc_year = 2026
                        else:
                            doc_year = 2026

                        if not doc_date or doc_date == "general":
                            doc_date = metadata_date_value(
                                pdf_doc.metadata,
                                pdf_doc.page_content[:1000],
                            )
                        if not doc_date or doc_date == "general":
                            doc_date = "2026-06-12"

                        pdf_doc.metadata["document_year"] = doc_year
                        pdf_doc.metadata["document_date"] = doc_date
                        pdf_doc.metadata["crawl_timestamp"] = pdf_doc.metadata.get("crawl_timestamp") or utc_timestamp()
                        pdf_doc.metadata["scope"] = "official"
                        pdf_doc.metadata["status"] = "active"
                        pdf_doc.metadata["deleted"] = False
                        pdf_doc.metadata["section_title"] = pdf_doc.metadata.get(
                            "section_title",
                            detect_section_title(pdf_doc.page_content),
                        )

                    documents.extend(pdf_docs)

                    print(
                        f"[WEBSITE DOC] Saved and Loaded content: {url} "
                        f"| file: {pdf_filename} | items: {len(pdf_docs)}"
                    )

                    if job_id:
                        is_pdf = is_pdf_url(url) or is_pdf_response
                        with crawl_jobs_lock:
                            job = crawl_jobs.get(job_id)
                            if job:
                                if is_pdf:
                                    job["pdfs_processed"] += 1
                                else:
                                    job["documents_processed"] += 1
                                job["current_stage"] = "completed"

                except Exception as e:
                    err_msg = f"Document extraction failed on {url}: {str(e)}"
                    if job_id:
                        update_crawl_job(job_id, errors=err_msg)
                    failed_pages.append(
                        {
                            "url": url,
                            "error": f"PDF extraction failed: {str(e)}",
                        }
                    )

                time.sleep(delay_seconds)
                continue

            # -----------------------------------------------------------------
            # HTML PAGE INGESTION
            # -----------------------------------------------------------------
            if not is_html_response(content_type, response.text):
                continue

            html = response.text

            # Checkpoint: before extracting links/content
            try:
                wait_if_paused(job_id)
                check_crawl_cancelled(job_id)
            except ValueError as e:
                print(f"[WEBSITE] Cancelled: {e}")
                break

            if should_skip_current_page(job_id):
                if job_id:
                    update_crawl_job(job_id, skipped_urls=url)
                continue

            main_text = extract_clean_text_from_website_html(html, url)
            links_text = extract_visible_links_text_from_html(html, url, base_domain)

            combined_text = ""

            if main_text:
                combined_text += main_text.strip()

            if links_text:
                combined_text += "\n\n" + links_text.strip()

            combined_text = clean_loaded_text(combined_text)

            # Checkpoint: before ingesting page text
            try:
                wait_if_paused(job_id)
                check_crawl_cancelled(job_id)
            except ValueError as e:
                print(f"[WEBSITE] Cancelled: {e}")
                break

            if should_skip_current_page(job_id):
                if job_id:
                    update_crawl_job(job_id, skipped_urls=url)
                continue

            if job_id:
                update_crawl_job(job_id, current_stage="ingesting")

            if word_count(combined_text) >= MIN_CHUNK_WORDS:
                page_title = readable_website_source_name(url, detect_section_title(combined_text))
                
                doc_year = metadata_year_value(url, page_title, combined_text[:3000])
                if doc_year and doc_year != "general":
                    try:
                        doc_year = int(doc_year)
                    except ValueError:
                        doc_year = 2026
                else:
                    doc_year = 2026

                doc_date = metadata_date_value({"source_url": url}, combined_text[:1000])
                if not doc_date or doc_date == "general":
                    doc_date = "2026-06-12"

                documents.append(
                    Document(
                        page_content=combined_text,
                        metadata={
                            "filename": page_title,
                            "source_filename": page_title,
                            "source_url": url,
                            "title": page_title,
                            "found_on_url": "",
                            "crawl_base_url": start_url,
                            "page": 1,
                            "total_pages": 0,
                            "file_type": "website",
                            "section_title": detect_section_title(combined_text),
                            "is_toc": detect_toc(combined_text),
                            "source_type": "website",
                            "document_year": doc_year,
                            "document_date": doc_date,
                            "crawl_timestamp": utc_timestamp(),
                            "scope": "official",
                            "status": "active",
                            "deleted": False,
                        },
                    )
                )

            # Task 7: Process links as website_links (own chunk/doc)
            if links_text and word_count(links_text) >= MIN_CHUNK_WORDS:
                total_links_extracted += links_text.count("- ")
                links_title = readable_website_source_name(url, "Website Links")
                
                doc_year = metadata_year_value(url, links_title, links_text[:3000])
                if doc_year and doc_year != "general":
                    try:
                        doc_year = int(doc_year)
                    except ValueError:
                        doc_year = 2026
                else:
                    doc_year = 2026

                doc_date = metadata_date_value({"source_url": url}, links_text[:1000])
                if not doc_date or doc_date == "general":
                    doc_date = "2026-06-12"

                documents.append(
                    Document(
                        page_content=links_text,
                        metadata={
                            "filename": links_title,
                            "source_filename": links_title,
                            "source_url": url,
                            "title": links_title,
                            "crawl_base_url": start_url, # Task 8
                            "file_type": "website_links",
                            "source_type": "website_links", # Task 8
                            "document_year": doc_year,
                            "document_date": doc_date,
                            "crawl_timestamp": utc_timestamp(),
                            "section_title": "Website Links",
                            "scope": "official",
                            "status": "active",
                            "deleted": False,
                        },
                    )
                )

            if job_id:
                with crawl_jobs_lock:
                    job = crawl_jobs.get(job_id)
                    if job:
                        job["pages_processed"] += 1
                        job["current_stage"] = "completed"

            # -----------------------------------------------------------------
            # LINK DISCOVERY
            # -----------------------------------------------------------------
            soup = BeautifulSoup(html or "", "html.parser")

            for tag in soup.find_all("a", href=True):
                href = tag.get("href", "").strip()

                if not href:
                    continue

                raw_absolute_url = urljoin(url, href)
                absolute_url = normalize_url(
                    raw_absolute_url,
                    keep_query=is_document_url(raw_absolute_url),
                )

                if not is_allowed_crawl_url(absolute_url):
                    continue

                if same_domain_only and not is_same_domain(absolute_url, base_domain):
                    continue

                label = tag.get_text(" ", strip=True)

                if not label:
                    label = os.path.basename(urlparse(absolute_url).path) or absolute_url

                if is_document_url(absolute_url):
                    if include_pdfs and absolute_url not in pdf_links_recorded:
                        from crawl4ai_crawler import should_skip_url
                        if should_skip_url(absolute_url):
                            print(f"[WEBSITE] Skipping document link by pattern: {absolute_url}")
                            continue

                        pdf_links_recorded.add(absolute_url)
                        pdf_found_on_urls[absolute_url] = url

                        documents.append(
                            build_pdf_link_document(
                                pdf_url=absolute_url,
                                pdf_title=label,
                                found_on_url=url,
                                start_url=start_url,
                                pdf_index=len(pdf_links_recorded),
                            )
                        )
                        if job_id:
                            is_pdf = is_pdf_url(absolute_url)
                            with crawl_jobs_lock:
                                job = crawl_jobs.get(job_id)
                                if job:
                                    if is_pdf:
                                        job["pdfs_found"] += 1
                                    else:
                                        job["documents_found"] += 1

                    if (
                        include_pdfs
                        and absolute_url not in visited
                        and absolute_url not in queued
                        and pdfs_loaded < max_pdfs
                    ):
                        from crawl4ai_crawler import should_skip_url
                        if should_skip_url(absolute_url):
                            print(f"[WEBSITE] Skipping document queueing by pattern: {absolute_url}")
                            continue
                        pdf_found_on_urls.setdefault(absolute_url, url)
                        queue.append((absolute_url, depth + 1))
                        queued.add(absolute_url)

                else:
                    if absolute_url not in visited and absolute_url not in queued:
                        from crawl4ai_crawler import should_skip_url
                        if should_skip_url(absolute_url):
                            print(f"[WEBSITE] Skipping page queueing by pattern: {absolute_url}")
                            continue
                        queue.append((absolute_url, depth + 1))
                        queued.add(absolute_url)
                        if job_id:
                            with crawl_jobs_lock:
                                job = crawl_jobs.get(job_id)
                                if job:
                                    job["pages_found"] += 1

            time.sleep(delay_seconds)

        except Exception as e:
            err_msg = f"Crawl loop error on {url}: {str(e)}"
            if job_id:
                update_crawl_job(job_id, errors=err_msg)
            failed_pages.append(
                {
                    "url": url,
                    "error": str(e),
                }
            )

    total_docs = len(documents)

    if not documents:
        failed_summary = ""
        if failed_pages:
            failed_sample = failed_pages[:3]
            failed_summary = f" Failed pages: {failed_sample}."

        raise ValueError(
            "No readable website content found. "
            f"Visited {len(visited)} page(s), discovered {len(pdf_links_recorded)} document link(s), "
            f"and loaded {pdfs_loaded} PDF/document(s).{failed_summary}"
        )

    for i, doc in enumerate(documents, start=1):
        doc.metadata["page"] = doc.metadata.get("page", i)
        doc.metadata["total_pages"] = total_docs
        doc.metadata["source_filename"] = doc.metadata.get("source_filename", start_url)
        doc.metadata["filename"] = doc.metadata.get("filename", start_url)
        doc.metadata["source_url"] = doc.metadata.get("source_url", "")
        doc.metadata["found_on_url"] = doc.metadata.get("found_on_url", "")
        doc.metadata["crawl_base_url"] = doc.metadata.get("crawl_base_url", start_url)
        doc.metadata["scope"] = doc.metadata.get("scope", "official")
        doc.metadata["status"] = doc.metadata.get("status", "active")
        doc.metadata["deleted"] = doc.metadata.get("deleted", False)
        
        # Add Crawl RAG metadata compliance fields
        doc.metadata["url"] = doc.metadata.get("url", doc.metadata.get("source_url", ""))
        doc.metadata["title"] = doc.metadata.get("title", doc.metadata.get("filename", start_url))
        doc.metadata["domain"] = doc.metadata.get("domain", base_domain)
        doc.metadata["source_type"] = "website"
        doc.metadata["crawl_timestamp"] = doc.metadata.get("crawl_timestamp", time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()))
        doc.metadata["crawl_method"] = "legacy"

    print("[WEBSITE CRAWL SUMMARY]")
    print("Start URL          :", start_url)
    print("Visited URLs       :", len(visited))
    print("Documents loaded   :", len(documents))
    print("PDF links recorded :", len(pdf_links_recorded))
    print("PDFs loaded        :", pdfs_loaded)
    print("Files downloaded   :", docs_downloaded)
    print("Total Links found  :", total_links_extracted)
    print("Failed pages       :", len(failed_pages))

    if failed_pages:
        print("Failed samples     :", failed_pages[:5])

    return documents


# =============================================================================
# FILE ROUTING
# =============================================================================

def load_file_from_path(
    file_path: str,
    job_id: str | None = None,
    url: str = None,
    extra_metadata: dict | None = None,
) -> list[Document]:
    ext = os.path.splitext(file_path)[1].lower()
    filename = os.path.basename(file_path)

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}")

    if ext == ".pdf":
        with open(file_path, "rb") as f:
            docs = load_pdf_bytes(f.read(), filename, job_id=job_id, extra_metadata=extra_metadata)

    elif ext in {".png", ".jpg", ".jpeg", ".webp"}:
        with open(file_path, "rb") as f:
            docs = load_image_bytes(f.read(), filename, url=url)

    elif ext == ".docx":
        docs = load_docx_file(file_path)

    elif ext == ".txt":
        docs = load_txt_file(file_path)

    elif ext in {".html", ".htm"}:
        docs = load_html_file(file_path)

    elif ext == ".csv":
        docs = load_csv_file(file_path)

    elif ext in {".md", ".markdown"}:
        docs = load_markdown_file(file_path)

    elif ext == ".json":
        docs = load_json_file(file_path)

    elif ext in {".xlsx", ".xls"}:
        docs = load_xlsx_file(file_path)

    elif ext in {".sql", ".dump"}:
        docs = load_db_dump_file(file_path)

    else:
        docs = []

    for i, doc in enumerate(docs, start=1):
        doc.metadata["filename"] = filename
        doc.metadata["file_type"] = doc.metadata.get("file_type", ext.replace(".", ""))
        doc.metadata["page"] = doc.metadata.get("page", i)
        doc.metadata["total_pages"] = doc.metadata.get("total_pages", len(docs))
        doc.metadata["section_title"] = doc.metadata.get(
            "section_title",
            detect_section_title(doc.page_content),
        )
        if extra_metadata:
            doc.metadata.update({str(k): str(v) for k, v in extra_metadata.items() if v is not None})

    return docs


def load_file_from_bytes(
    file_bytes: bytes,
    filename: str,
    job_id: str | None = None,
    url: str = None,
    extra_metadata: dict | None = None,
) -> list[Document]:
    ext = os.path.splitext(filename)[1].lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}")

    if ext == ".pdf":
        return load_pdf_bytes(file_bytes, filename, job_id=job_id, extra_metadata=extra_metadata)

    if ext in {".png", ".jpg", ".jpeg", ".webp"}:
        return load_image_bytes(file_bytes, filename, url=url)

    temp_dir = "temp_uploads"
    os.makedirs(temp_dir, exist_ok=True)

    temp_path = os.path.join(temp_dir, f"{uuid.uuid4()}_{safe_id_text(filename)}")

    try:
        with open(temp_path, "wb") as f:
            f.write(file_bytes)

        documents = load_file_from_path(temp_path, job_id=job_id, url=url, extra_metadata=extra_metadata)

        for doc in documents:
            doc.metadata["filename"] = filename
            doc.metadata["section_title"] = doc.metadata.get(
                "section_title",
                detect_section_title(doc.page_content),
            )

        return documents

    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


# =============================================================================
# INGESTION
# =============================================================================

def ingest_documents(
    documents: list[Document],
    filename: str,
    session_id: str = None,
    user_id: str = None,
    department: str = "general",
    document_type: str = "general",
    year: str = "general",
    scope: str = "official",
    crawl_base_url: str = "",
    job_id: str | None = None,
    rebuild_bm25: bool = True,
) -> dict:
    if job_id:
        update_crawl_job(job_id, status="processing", current_stage="ingesting")
    scope = "personal" if str(scope).lower() == "personal" else "official"
    is_personal = scope == "personal"

    uploaded_by = "user" if is_personal else "admin"

    final_session_id = session_id or ("personal" if is_personal else "admin")
    final_user_id = user_id or ("personal" if is_personal else "admin")

    department = str(department or "general")
    document_type = str(document_type or "general")
    year = str(year or "general")

    safe_filename = safe_id_text(filename)
    safe_user_id = safe_id_text(final_user_id)
    safe_session_id = safe_id_text(final_session_id)

    all_chunks: list[str] = []
    all_embedding_texts: list[str] = []
    all_metadatas: list[dict] = []
    all_ids: list[str] = []

    total_text_chars = 0
    readable_pages = 0
    tables_extracted = 0
    ocr_pages_used = 0
    chunks_skipped = 0
    duplicates_skipped = 0
    candidate_chunks = 0

    seen_hashes: set[str] = set()
    current_run_hashes: set[str] = set()
    # Near-duplicate index for this run, bucketed by digit-signature so we only
    # compare SimHashes among chunks that share the exact same numbers.
    kept_simhashes_by_digits: dict[tuple, list[int]] = {}

    existing_hashes: set[str] = set()
    existing_lookup_field = "crawl_base_url" if crawl_base_url else "filename"
    existing_lookup_value = crawl_base_url or filename

    try:
        existing = collection.get(
            where={existing_lookup_field: {"$eq": existing_lookup_value}},
            include=["metadatas"],
        )

        existing_hashes = {
            str(meta.get("text_hash"))
            for meta in existing.get("metadatas", [])
            if meta and meta.get("text_hash")
        }

    except Exception as e:
        print(f"Warning: failed to read existing chunk hashes for dedupe: {e}")

    print(f"[DEBUG INGEST] existing_lookup_value={existing_lookup_value!r} existing_hashes count={len(existing_hashes)}")

    website_pdf_count = 0
    website_pdf_link_count = 0
    website_html_count = 0
    website_links_count = 0
    source_urls_seen: set[str] = set()

    chunk_in_file_index = 0

    for doc in documents:
        if job_id:
            wait_if_paused(job_id)
            check_crawl_cancelled(job_id)

        text = clean_loaded_text(doc.page_content)

        if not text:
            continue

        readable_pages += 1
        total_text_chars += len(text)

        tables_extracted += int(doc.metadata.get("tables_extracted", 0) or 0)

        if doc.metadata.get("ocr_used") is True:
            ocr_pages_used += 1

        source_url = str(doc.metadata.get("source_url", "") or "")
        doc_filename = str(doc.metadata.get("filename", filename) or filename)
        doc_source_filename = str(doc.metadata.get("source_filename", doc_filename) or doc_filename)

        if source_url:
            source_urls_seen.add(source_url)

        file_type = str(doc.metadata.get("file_type", "") or "")
        source_type = doc.metadata.get("source_type")
        if not source_type:
            if crawl_base_url or source_url.startswith(("http://", "https://")):
                source_type = "website_document" if file_type == "website_document" else "website_page"
            else:
                source_type = "uploaded_document"
            if file_type in {"website_pdf", "pdf"} and source_url.startswith(("http://", "https://")):
                source_type = "website_pdf"
            if file_type == "website_document":
                source_type = "website_document"
            if document_type == "general" and file_type == "website":
                source_type = "website_page"

        if source_type == "website_pdf" or file_type == "website_pdf":
            website_pdf_count += 1
        elif file_type == "website_pdf_link":
            website_pdf_link_count += 1
        elif file_type == "website":
            website_html_count += 1
        elif file_type == "website_links":
            website_links_count += 1

        page_number = doc.metadata.get("page", 1)
        section_title = doc.metadata.get("section_title") or detect_section_title(text)
        doc_is_toc = bool(doc.metadata.get("is_toc", False))

        if job_id:
            wait_if_paused(job_id)
            check_crawl_cancelled(job_id)
            update_crawl_job(
                job_id,
                current_url=source_url or doc_filename,
                current_type="pdf" if file_type in {"pdf", "website_pdf"} else ("document" if file_type == "website_document" else "page"),
                current_stage="chunking",
                status="chunking"
            )

        chunks = chunk_text(text)
        # Split any chunk that would overflow the embedding model's token budget
        # so its full content is embedded (not silently truncated by the encoder).
        chunks = split_chunks_for_embedding(chunks)
        candidate_chunks += len(chunks)
        char_search_start = 0

        for chunk in chunks:
            chunk_index = chunk_in_file_index
            chunk_in_file_index += 1

            # chunk_text() already enforced MIN_CHUNK_WORDS (with the
            # high-value-short-chunk exemption). Re-apply the same exemption here
            # rather than a blanket word-count drop, so short contact/fee/role
            # records that legitimately passed chunking are not discarded now.
            if len(chunk.split()) < MIN_CHUNK_WORDS and not is_valuable_short_chunk(chunk):
                chunks_skipped += 1
                continue

            if source_type == "website_links" and chunk_index == 0:
                print("\n[DEBUG] Sample website_links chunk:")
                print(f"Text: {chunk[:200]}...")
                print(f"Metadata: source_type={source_type}, filename={filename}")

            doc_id = compute_doc_id(doc_filename, source_url)
            page_label = str(doc.metadata.get("page_label", page_number))
            heading_path = build_heading_path(section_title, doc.metadata)
            chunk_type = detect_chunk_type(chunk)
            table_title = detect_table_title(chunk, section_title)
            char_start, char_end = find_chunk_char_offsets(text, chunk, char_search_start)
            char_search_start = char_end

            source_filename_val = doc_source_filename or doc_filename or filename or "unknown"
            extracted_title_val = doc.metadata.get("title") or doc.metadata.get("pdf_title") or source_filename_val
            current_heading_val = doc.metadata.get("heading") or section_title or ""
            section_name_val = doc.metadata.get("section") or section_title or ""
            doc_category_val = doc.metadata.get("category") or doc.metadata.get("department") or department or "general"

            file_extension_val = doc.metadata.get("doc_type")
            if not file_extension_val:
                if "." in source_filename_val:
                    file_extension_val = source_filename_val.split(".")[-1].lower()
                else:
                    file_extension_val = doc.metadata.get("file_type") or "unknown"

            # Document year/date must be HONEST. Keep the value the loader
            # extracted, otherwise derive a year only from trusted identifier
            # fields (filename / title / URL / section). If still unknown, store
            # "general"/"" — never a fabricated current-year stamp. freshness.py
            # treats an unknown year as "no recency boost"; a fake "2026" would
            # make every undated old document masquerade as the freshest and
            # corrupt freshness/authority conflict resolution (e.g. "current
            # principal", "latest fee structure"). A crawl timestamp is likewise
            # NOT a publication date, so it is never used as the fallback here.
            doc_year = doc.metadata.get("document_year")
            try:
                doc_year = (
                    int(doc_year)
                    if str(doc_year).strip().lower() not in ("", "general", "none", "null")
                    else None
                )
            except (TypeError, ValueError):
                doc_year = None
            if not doc_year:
                derived_year = extract_year_from_text(
                    doc_filename,
                    doc.metadata.get("title"),
                    source_url,
                    section_title,
                )
                doc_year = derived_year if derived_year else "general"

            doc_date = doc.metadata.get("document_date")
            if not doc_date or str(doc_date).strip().lower() in ("general", "none", "null"):
                doc_date = ""

            metadata_dict = {
                **{k: normalize_metadata_value(v) for k, v in doc.metadata.items()},
                "document_year": doc_year,
                "document_date": doc_date,
                "filename": source_filename_val,
                "title": extracted_title_val,
                "heading": current_heading_val,
                "page": page_number or 0,
                "section": section_name_val,
                "category": doc_category_val,
                "doc_type": file_extension_val,
                "page_range": normalize_metadata_value(
                    doc.metadata.get("page_range", str(page_number)),
                    str(page_number),
                ),
                "total_pages": normalize_metadata_value(
                    doc.metadata.get("total_pages", len(documents)),
                    len(documents),
                ),
                "file_type": file_extension_val,
                "section_title": current_heading_val,
                "source_url": normalize_metadata_value(source_url, ""),
                "source": normalize_metadata_value(source_url or doc_filename, "unknown"),
                "found_on_url": normalize_metadata_value(
                    doc.metadata.get("found_on_url", ""),
                    "",
                ),
                "source_pdf_filename": normalize_metadata_value(
                    doc.metadata.get("source_pdf_filename", ""),
                    "",
                ),
                "pdf_title": normalize_metadata_value(
                    doc.metadata.get("pdf_title", ""),
                    "",
                ),
                "doc_id": doc_id,
                "page_label": page_label,
                "heading_path": heading_path,
                "chunk_type": chunk_type,
                "content_type": chunk_type,
                "table_title": table_title,
                "text_chars": len(chunk),
                "char_start": char_start,
                "char_end": char_end,
                "chunk_index": chunk_index,
                "ingested_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
                "deleted": False,
                "status": "active",
                "scope": scope,
                "uploaded_by": uploaded_by,
                "user_id": final_user_id,
                "session_id": final_session_id,
                "department": department,
                "document_type": document_type,
                "year": year,
                "source_type": source_type,
                "crawl_base_url": crawl_base_url or filename,
                "word_count": len(chunk.split()),
                "is_toc": doc_is_toc or detect_toc(chunk),
            }

            for key, val in list(metadata_dict.items()):
                if val is None:
                    if key in ("page", "chunk_index", "word_count", "char_count", "text_chars", "total_pages"):
                        metadata_dict[key] = 0
                    elif key in ("deleted", "is_toc", "ocr_used", "tables_extracted"):
                        metadata_dict[key] = False
                    elif key in ("category", "document_type", "department", "year"):
                        metadata_dict[key] = "general"
                    elif key in ("doc_type", "file_type"):
                        metadata_dict[key] = "unknown"
                    else:
                        metadata_dict[key] = ""

            metadata = normalize_metadata(metadata_dict)

            import json
            # Chunk identity = content + STABLE source location only.
            #
            # Previously the hash mixed in nearly all metadata, so re-ingesting an
            # unchanged document after an admin edited its department / document_type
            # / year (or after total_pages shifted between crawls, or after the
            # honest-date backfill) changed every chunk's hash -> new id -> needless
            # re-embedding AND orphan-deletion churn. Restricting the hash to the
            # content plus its source identity (filename / source_url / page /
            # chunk_index) makes re-ingestion a true no-op when only classification
            # metadata changed, while still keeping identical text from DIFFERENT
            # sources distinct (different filename/source_url -> different hash), so
            # distinct factual records are never collapsed.
            identity = {
                "filename": metadata.get("filename", ""),
                "source_url": metadata.get("source_url", ""),
                "page": metadata.get("page", 0),
                "chunk_index": metadata.get("chunk_index", 0),
            }
            combined_text = chunk + json.dumps(identity, sort_keys=True)
            text_hash = hashlib.sha256(combined_text.encode("utf-8")).hexdigest()[:24]
            metadata["text_hash"] = text_hash

            current_run_hashes.add(text_hash)

            if text_hash in seen_hashes or text_hash in existing_hashes:
                duplicates_skipped += 1
                continue

            seen_hashes.add(text_hash)

            # Near-duplicate guard (within this ingest run): skip a chunk that is
            # ~identical to one already kept this run. The digit-signature bucket
            # ensures we only ever merge chunks with the SAME numbers, so distinct
            # numeric data is never collapsed. Exact re-ingest duplicates are
            # already handled above, so this only fires for genuinely new chunks.
            digit_sig = _digit_signature(chunk)
            chunk_simhash = _simhash(chunk)
            bucket = kept_simhashes_by_digits.setdefault(digit_sig, [])
            if any(_hamming(chunk_simhash, kept) <= SIMHASH_NEAR_DUP_MAX_HAMMING for kept in bucket[-500:]):
                duplicates_skipped += 1
                continue
            bucket.append(chunk_simhash)

            embedding_text = (
                f"Title: {metadata.get('title', '')}\n"
                f"Source: {metadata.get('filename', '')}\n"
                f"Section: {metadata.get('heading', '')}\n\n"
                f"{chunk}"
            )

            chunk_id = (
                f"personal_{safe_user_id}_{safe_session_id}_{safe_filename}"
                f"_p{str(page_number)}_c{chunk_index}_{text_hash}"
                if is_personal
                else f"official_{safe_filename}_p{str(page_number)}_c{chunk_index}_{text_hash}"
            )

            # KNOWLEDGE HIERARCHY: stamp authority metadata (document_type,
            # category, priority_level, authority_score, hostel_type, display_name,
            # version) so Tier 1 canonical sources are distinguishable at retrieval
            # time, plus spec-named aliases. Existing chunks are handled by
            # scripts/backfill_authority_metadata.py.
            try:
                from rag.authority import classify_document
                metadata.update(classify_document(metadata))
            except Exception as _authority_err:
                print(f"[INGEST] authority classify skipped: {_authority_err}")
            metadata.setdefault("section_heading", metadata.get("section_title", ""))
            metadata.setdefault("page_number", metadata.get("page", page_number))

            all_chunks.append(chunk)
            all_embedding_texts.append(embedding_text)
            all_metadatas.append(metadata)
            all_ids.append(chunk_id)

            if chunk_index < 5:
                print(
                    "[INGEST DEBUG]",
                    "file=",
                    filename,
                    "page=",
                    page_number,
                    "section=",
                    section_title,
                    "source_url=",
                    source_url,
                    "file_type=",
                    file_type,
                    "chunk=",
                    chunk_index,
                    "preview=",
                    embedding_text[:180],
                )

    # When every produced chunk was already stored and unchanged (duplicates_skipped),
    # this is a successful no-op re-ingest, not a failure. Fall through so orphan
    # deletion can still prune pages that genuinely disappeared, then return zero new
    # chunks. Only a truly empty extraction (no readable text at all) is an error.
    if not all_chunks and not duplicates_skipped:
        raise ValueError(f"No readable text found in '{filename}'")

    numeric_page_counts = [
        int(doc.metadata["total_pages"])
        for doc in documents
        if str(doc.metadata.get("total_pages", "")).isdigit()
    ]

    total_pages = max(numeric_page_counts) if numeric_page_counts else len(documents)

    print("[INGEST SUMMARY]", "file=", filename)
    print("[INGEST SUMMARY]", "pages_processed=", total_pages)
    print("[INGEST SUMMARY]", "readable_pages=", readable_pages)
    print("[INGEST SUMMARY]", "text_extracted_chars=", total_text_chars)
    print("[INGEST SUMMARY]", "chunks_stored will be=", len(all_chunks))

    # One-line explanation: Check which chunk IDs already exist in ChromaDB to make the ingestion resumable and avoid redundant embedding calls.
    existing_ids = set()
    for i in range(0, len(all_ids), 1000):
        batch_ids_to_check = all_ids[i : i + 1000]
        try:
            res = collection.get(ids=batch_ids_to_check, include=[])
            existing_ids.update(res.get("ids", []))
        except Exception as e:
            print(f"Warning: failed to check existing IDs batch: {e}")

    print(f"[INGEST] Found {len(existing_ids)} chunks already stored in ChromaDB out of {len(all_ids)} total chunks.")

    # One-line explanation: Embed and store chunks in batches to support incremental updates and log progress dynamically.
    chunks_stored = 0
    batch_size = 128
    for i in range(0, len(all_chunks), batch_size):
        batch_chunks = all_chunks[i : i + batch_size]
        batch_texts = all_embedding_texts[i : i + batch_size]
        batch_metadatas = all_metadatas[i : i + batch_size]
        batch_ids = all_ids[i : i + batch_size]

        # Identify which chunks in this batch are NOT already in ChromaDB
        indices_to_embed = [idx for idx, cid in enumerate(batch_ids) if cid not in existing_ids]

        if indices_to_embed:
            embed_texts = [batch_texts[idx] for idx in indices_to_embed]
            embed_chunks = [batch_chunks[idx] for idx in indices_to_embed]
            embed_metadatas = [batch_metadatas[idx] for idx in indices_to_embed]
            embed_ids = [batch_ids[idx] for idx in indices_to_embed]

            if job_id:
                wait_if_paused(job_id)
                check_crawl_cancelled(job_id)
                update_crawl_job(
                    job_id,
                    current_stage="embedding",
                    status="embedding",
                    embeddings_generated=chunks_stored + len(embed_ids)
                )

            # Generate embeddings for the missing chunks
            batch_embeddings = encode_texts(embed_texts, batch_size=len(embed_texts))

            # Store them in ChromaDB
            stored = add_chunks(
                chunks=embed_chunks,
                filename=filename,
                embeddings=batch_embeddings,
                metadatas=embed_metadatas,
                ids=embed_ids,
            )
            chunks_stored += stored
            print(f"[INGEST] Embedded and stored {stored} new chunks (progress: {i + len(batch_chunks)}/{len(all_chunks)})")
        else:
            chunks_stored += len(batch_ids)
            print(f"[INGEST] Skipped {len(batch_ids)} chunks (already stored) (progress: {i + len(batch_chunks)}/{len(all_chunks)})")

    if job_id:
        update_crawl_job(job_id, chunks_created=chunks_stored)

    # One-line explanation: Delete orphaned chunks (old chunks that are no longer present in the newly ingested document/page).
    #
    # SAFETY (partial-crawl data loss): a website crawl calls ingest_documents
    # ONCE with crawl_base_url set, so existing_lookup_value is the WHOLE site.
    # A recrawl that returns fewer pages — lower max_pages, reduced max_depth,
    # transient network failures, skipped pages, or an early stop — must NOT be
    # treated as "those pages were removed from the source". We therefore only
    # prune stale chunks belonging to URLs we ACTUALLY revisited this run
    # (source_urls_seen); chunks for pages not fetched this run are preserved.
    # Uploads (no crawl_base_url) keep whole-file scope so a modified re-upload
    # still replaces its own old chunks. A cancelled job skips pruning entirely.
    is_crawl_scope = bool(crawl_base_url)
    crawl_was_cancelled = bool(job_id) and should_cancel(job_id)
    if not is_personal and not crawl_was_cancelled:
        try:
            existing_doc_res = collection.get(
                where={existing_lookup_field: {"$eq": existing_lookup_value}},
                include=["metadatas"]
            )
            existing_doc_ids = existing_doc_res.get("ids", [])
            existing_doc_metas = existing_doc_res.get("metadatas", []) or []
            all_ids_set = set(all_ids)
            # Keep a stored chunk if it was (re)generated this run (id in all_ids_set)
            # OR its exact text is still present in this run (text_hash in
            # current_run_hashes — covers unchanged chunks skipped as duplicates).
            # Only delete chunks whose content is genuinely gone from the source.
            ids_to_delete = []
            for idx, existing_doc_id in enumerate(existing_doc_ids):
                meta = existing_doc_metas[idx] if idx < len(existing_doc_metas) else {}
                existing_text_hash = str((meta or {}).get("text_hash", "") or "")
                if existing_doc_id in all_ids_set or (
                    existing_text_hash and existing_text_hash in current_run_hashes
                ):
                    continue
                # For a crawl, only prune chunks whose own page was revisited this
                # run. A chunk with a source_url we never fetched (or no source_url
                # to verify) is left untouched — a partial crawl cannot delete it.
                if is_crawl_scope:
                    existing_source_url = str((meta or {}).get("source_url", "") or "")
                    if existing_source_url not in source_urls_seen:
                        continue
                ids_to_delete.append(existing_doc_id)
            if ids_to_delete:
                print(f"[INGEST] Deleting {len(ids_to_delete)} orphaned chunks for {existing_lookup_value}...")
                # Delete in batches to prevent exceeding max batch size in deletes as well
                max_delete_batch = 1000
                for d_idx in range(0, len(ids_to_delete), max_delete_batch):
                    collection.delete(ids=ids_to_delete[d_idx:d_idx+max_delete_batch])
        except Exception as e:
            print(f"Warning: failed to delete orphaned chunks: {e}")

    # §8: Invalidate response and retrieval caches after ingestion
    try:
        from rag.cache import invalidate_on_ingestion
        invalidate_on_ingestion()
        print("[INGEST] Cache invalidated (Layers 1 & 2)")
    except Exception as e:
        print(f"[INGEST] Warning: Cache invalidation failed: {e}")

    # Rebuild the BM25 keyword index so freshly ingested chunks are immediately
    # searchable by the lexical/BM25 + special-list retrieval paths. Previously
    # this happened only on server startup or in ingest_folder, so /upload and
    # /crawl content was invisible to keyword retrieval until a restart. Bulk
    # callers (ingest_folder) pass rebuild_bm25=False and rebuild once at the end.
    if rebuild_bm25:
        try:
            rebuild_bm25_index()
            print("[INGEST] BM25 index rebuilt")
        except Exception as e:
            print(f"[INGEST] Warning: BM25 rebuild failed: {e}")

    try:
        collection_count = collection.count()
    except Exception:
        collection_count = None

    print("[INGEST SUMMARY]", "chunks_stored=", chunks_stored)
    print("[INGEST SUMMARY]", "collection_count=", collection_count)

    ext = os.path.splitext(filename)[1].replace(".", "").upper()
    file_type_label = ext or str(document_type or "UNKNOWN").upper()

    return {
        "file": filename,
        "type": file_type_label,
        "scope": scope,
        "uploaded_by": uploaded_by,
        "department": department,
        "document_type": document_type,
        "year": year,
        "pages_processed": total_pages,
        "readable_pages": readable_pages,
        "text_extracted_chars": total_text_chars,
        "chunks_created": candidate_chunks,
        "chunks_skipped": chunks_skipped,
        "duplicates_skipped": duplicates_skipped,
        "chunks_stored": chunks_stored,
        "chunk_size_chars": MAX_CHARS_PER_CHUNK,
        "chunk_overlap_chars": CHUNK_OVERLAP,
        "tables_extracted": tables_extracted,
        "ocr_pages_used": ocr_pages_used,
        "collection_count": collection_count,
        "website_html_docs": website_html_count,
        "website_pdf_docs": website_pdf_count,
        "website_pdf_link_docs": website_pdf_link_count,
        "website_links_docs": website_links_count,
        "unique_source_urls": len(source_urls_seen),
        "include_pdfs": None,
        "max_pdfs": None,
        "status": "Ready for RAG search",
    }


# =============================================================================
# DEBUG SEARCH
# =============================================================================

def debug_search_ingested_text(term: str):
    def norm(value: str) -> str:
        value = (value or "").lower()
        value = re.sub(r"[^a-z0-9\s]", " ", value)

        return re.sub(r"\s+", " ", value).strip()

    result = collection.get(include=["documents", "metadatas"])
    docs = result.get("documents", [])
    metas = result.get("metadatas", [])

    print("Total chunks:", len(docs))
    print("Searching for:", term)

    found = 0

    for doc, meta in zip(docs, metas):
        if norm(term) in norm(doc):
            found += 1

            print("\nFOUND")
            print("File       :", meta.get("filename"))
            print("Page       :", meta.get("page"))
            print("Section    :", meta.get("section_title"))
            print("File type  :", meta.get("file_type"))
            print("Source URL :", meta.get("source_url"))
            print("PDF title  :", meta.get("pdf_title"))
            print("PDF file   :", meta.get("source_pdf_filename"))
            print("Chunk      :", meta.get("chunk_index"))
            print("Preview    :")
            print(doc[:800])

    print("\nMatches found:", found)


# =============================================================================
# PUBLIC ENTRY POINTS
# =============================================================================

def ingest_file_bytes(
    file_bytes: bytes,
    filename: str,
    session_id: str = None,
    user_id: str = None,
    department: str = "general",
    document_type: str = "general",
    year: str = "general",
    scope: str = "official",
) -> dict:
    documents = load_file_from_bytes(file_bytes, filename)

    return ingest_documents(
        documents=documents,
        filename=filename,
        session_id=session_id,
        user_id=user_id,
        department=department,
        document_type=document_type,
        year=year,
        scope=scope,
    )


def ingest_file_path(
    file_path: str,
    session_id: str = None,
    user_id: str = None,
    department: str = "general",
    document_type: str = "general",
    year: str = "general",
    scope: str = "official",
    rebuild_bm25: bool = True,
) -> dict:
    filename = os.path.basename(file_path)
    documents = load_file_from_path(file_path)

    return ingest_documents(
        documents=documents,
        filename=filename,
        session_id=session_id,
        user_id=user_id,
        department=department,
        document_type=document_type,
        year=year,
        scope=scope,
        rebuild_bm25=rebuild_bm25,
    )


def ingest_website(
    url: str,
    session_id: str = None,
    user_id: str = None,
    department: str = "general",
    document_type: str = "website",
    year: str = "general",
    scope: str = "official",
    max_pages: int = 50,
    delay_seconds: float = 0.5,
    include_pdfs: bool = True,
    max_pdfs: int = 200,
    same_domain_only: bool = True,
    job_id: str | None = None,
    max_depth: int = 3,
) -> dict:
    cleaned_url = normalize_url(url)

    # Hard domain lock (anti-SSRF + anti-off-domain crawl). Fail fast with a clear
    # message instead of silently crawling nothing when the loop drops every URL.
    if is_private_ip(cleaned_url) or not is_domain_allowed(cleaned_url):
        allowed = ", ".join(sorted(ALLOWED_CRAWL_DOMAINS)) or "(none configured)"
        raise ValueError(
            f"Refusing to crawl {cleaned_url!r}: host is outside the allowed college "
            f"domain(s): {allowed}. Set ALLOWED_CRAWL_DOMAINS to permit another domain."
        )

    if job_id:
        check_crawl_cancelled(job_id)

    documents = load_website(
        start_url=cleaned_url,
        max_pages=max_pages,
        delay_seconds=delay_seconds,
        include_pdfs=include_pdfs,
        max_pdfs=max_pdfs,
        same_domain_only=same_domain_only,
        job_id=job_id,
        max_depth=max_depth,
    )

    if job_id:
        check_crawl_cancelled(job_id)

    result = ingest_documents(
        documents=documents,
        filename=cleaned_url,
        session_id=session_id,
        user_id=user_id,
        department=department,
        document_type=document_type,
        year=year,
        scope=scope,
        crawl_base_url=cleaned_url,
        job_id=job_id,
    )

    print(f"\n[WEBSITE CRAWL COMPLETED]")
    print(f"Chunks stored: {result.get('chunks_stored', 0)}")
    print(f"Link documents processed: {result.get('website_links_docs', 0)}")

    result["include_pdfs"] = include_pdfs
    result["max_pdfs"] = max_pdfs
    result["success"] = True
    return result


def ingest_from_url(
    url: str,
    collection=None,
    embedder=None,
    max_pages: int = 50,
) -> dict:
    """
    Backward-compatible wrapper.
    """
    return ingest_website(
        url=url,
        max_pages=max_pages,
        scope="official",
        include_pdfs=True,
    )


def ingest_folder(
    folder_path: str = "data",
    department: str = "general",
    document_type: str = "general",
    year: str = "general",
    scope: str = "official",
) -> list[dict]:
    if not os.path.exists(folder_path):
        raise FileNotFoundError(f"Folder not found: {folder_path}")

    results: list[dict] = []

    for root, _dirs, files in os.walk(folder_path):
        for filename in files:
            ext = os.path.splitext(filename)[1].lower()

            if ext not in SUPPORTED_EXTENSIONS:
                continue

            file_path = os.path.join(root, filename)

            try:
                result = ingest_file_path(
                    file_path,
                    session_id=None,
                    user_id=None,
                    department=department,
                    document_type=document_type,
                    year=year,
                    scope=scope,
                    # Defer BM25 rebuild to the single call after the loop so a
                    # large folder ingest doesn't rebuild the index per file.
                    rebuild_bm25=False,
                )

                results.append(result)
                print(f"\nIngestion complete: {result['file']}")

            except Exception as e:
                results.append(
                    {
                        "file": filename,
                        "status": "Failed",
                        "error": str(e),
                    }
                )

                print(f"\nIngestion failed: {filename} — {e}")

    # Rebuild BM25 Index after ingestion
    rebuild_bm25_index()

    return results


if __name__ == "__main__":
    # Normal folder ingestion
    # ingest_folder("data")

    # Website ingestion test
    result = ingest_website(
        url="https://anthonys.ac.in/",
        max_pages=50,
        include_pdfs=True,
        max_pdfs=200,
        scope="official",
        department="general",
        document_type="website",
        year="general",
    )

    print(result)
